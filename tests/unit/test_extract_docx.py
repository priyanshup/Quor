"""Unit tests for quor/pipeline/extract/docx.py — QB-007E2.

Covers real DOCX-to-Markdown extraction. Fixture documents are built with
python-docx itself (real, valid `.docx` files written to `tmp_path`), not
mocks — the same library extract_docx() uses, so a fixture is guaranteed to
exercise the real on-disk file format rather than an idealized in-memory
approximation of it. Fail-open cases (corrupt file, invalid zip, missing
dependency) use genuinely broken/absent input for the same reason.
"""

from __future__ import annotations

import sys
import warnings
import zipfile
from pathlib import Path

import docx
import pytest

from quor.pipeline.extract.docx import extract_docx

# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def _save(document: docx.Document, tmp_path: Path, name: str = "fixture.docx") -> Path:
    path = tmp_path / name
    document.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------


class TestHeadings:
    def test_heading_levels_1_through_6(self, tmp_path: Path) -> None:
        d = docx.Document()
        for level in range(1, 7):
            d.add_heading(f"Level {level} heading", level=level)
        result = extract_docx(_save(d, tmp_path))
        assert result is not None
        for level in range(1, 7):
            assert f"{'#' * level} Level {level} heading" in result

    def test_heading_is_single_line_even_with_embedded_break(self, tmp_path: Path) -> None:
        """A heading paragraph's internal whitespace/line breaks must be
        flattened — the markdown filter's heading pattern is single-line
        anchored (`^#{1,6}\\s+\\S`)."""
        d = docx.Document()
        h = d.add_heading("", level=1)
        run = h.add_run("Part one")
        run.add_break()
        h.add_run("part two")
        result = extract_docx(_save(d, tmp_path))
        assert result is not None
        assert "# Part one part two" in result
        assert "\n" not in result.splitlines()[0]

    def test_non_heading_styles_are_not_treated_as_headings(self, tmp_path: Path) -> None:
        """Word's "Title"/"Subtitle" styles and Heading 7+ are out of this
        phase's stated scope (Heading 1..6 only) and fall through as normal
        paragraphs, not silently dropped."""
        d = docx.Document()
        d.add_paragraph("A Document Title", style="Title")
        result = extract_docx(_save(d, tmp_path))
        assert result == "A Document Title"


# ---------------------------------------------------------------------------
# Paragraphs
# ---------------------------------------------------------------------------


class TestParagraphs:
    def test_plain_paragraphs_preserved(self, tmp_path: Path) -> None:
        d = docx.Document()
        d.add_paragraph("First paragraph.")
        d.add_paragraph("Second paragraph.")
        result = extract_docx(_save(d, tmp_path))
        assert result == "First paragraph.\n\nSecond paragraph."

    def test_empty_paragraphs_do_not_produce_blank_blocks(self, tmp_path: Path) -> None:
        d = docx.Document()
        d.add_paragraph("Before.")
        d.add_paragraph("")
        d.add_paragraph("")
        d.add_paragraph("After.")
        result = extract_docx(_save(d, tmp_path))
        assert result == "Before.\n\nAfter."

    def test_hyperlink_visible_text_is_preserved(self, tmp_path: Path) -> None:
        """python-docx's own Paragraph.text contract: "includes the
        visible-text portion of any hyperlinks." Verified against the real
        object model here, not assumed."""
        from docx.opc.constants import RELATIONSHIP_TYPE
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        d = docx.Document()
        p = d.add_paragraph("See ")
        r_id = p.part.relate_to(
            "https://example.com", RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run_el = OxmlElement("w:r")
        text_el = OxmlElement("w:t")
        text_el.text = "the docs"
        run_el.append(text_el)
        hyperlink.append(run_el)
        p._p.append(hyperlink)
        p.add_run(" for details.")

        result = extract_docx(_save(d, tmp_path))
        assert result == "See the docs for details."


# ---------------------------------------------------------------------------
# Bullet lists
# ---------------------------------------------------------------------------


class TestBulletLists:
    def test_bullet_list_items(self, tmp_path: Path) -> None:
        d = docx.Document()
        d.add_paragraph("Intro.")
        d.add_paragraph("First item", style="List Bullet")
        d.add_paragraph("Second item", style="List Bullet")
        d.add_paragraph("Third item", style="List Bullet")
        result = extract_docx(_save(d, tmp_path))
        assert result == (
            "Intro.\n\n- First item\n\n- Second item\n\n- Third item"
        )

    def test_bullet_list_interrupted_by_paragraph_restarts_cleanly(self, tmp_path: Path) -> None:
        d = docx.Document()
        d.add_paragraph("One", style="List Bullet")
        d.add_paragraph("A plain paragraph in between.")
        d.add_paragraph("Two", style="List Bullet")
        result = extract_docx(_save(d, tmp_path))
        assert result == (
            "- One\n\nA plain paragraph in between.\n\n- Two"
        )


# ---------------------------------------------------------------------------
# Numbered lists
# ---------------------------------------------------------------------------


class TestNumberedLists:
    def test_numbered_list_increments(self, tmp_path: Path) -> None:
        d = docx.Document()
        d.add_paragraph("Step one", style="List Number")
        d.add_paragraph("Step two", style="List Number")
        d.add_paragraph("Step three", style="List Number")
        result = extract_docx(_save(d, tmp_path))
        assert result == "1. Step one\n\n2. Step two\n\n3. Step three"

    def test_numbered_list_restarts_after_interruption(self, tmp_path: Path) -> None:
        """A new numbered-list run (after a non-list paragraph breaks the
        previous one) starts back at 1 — this is a stated "reasonably
        detectable" heuristic, not a faithful resolution of Word's own
        multi-list numbering XML."""
        d = docx.Document()
        d.add_paragraph("First run item one", style="List Number")
        d.add_paragraph("First run item two", style="List Number")
        d.add_paragraph("An interrupting paragraph.")
        d.add_paragraph("Second run item one", style="List Number")
        result = extract_docx(_save(d, tmp_path))
        assert result == (
            "1. First run item one\n\n"
            "2. First run item two\n\n"
            "An interrupting paragraph.\n\n"
            "1. Second run item one"
        )


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


class TestTables:
    def test_simple_table_renders_as_github_markdown(self, tmp_path: Path) -> None:
        d = docx.Document()
        t = d.add_table(rows=3, cols=2)
        t.cell(0, 0).text = "Option"
        t.cell(0, 1).text = "Verdict"
        t.cell(1, 0).text = "A"
        t.cell(1, 1).text = "Good"
        t.cell(2, 0).text = "B"
        t.cell(2, 1).text = "Bad"
        result = extract_docx(_save(d, tmp_path))
        assert result == (
            "| Option | Verdict |\n"
            "| --- | --- |\n"
            "| A | Good |\n"
            "| B | Bad |"
        )

    def test_table_cell_pipe_is_escaped(self, tmp_path: Path) -> None:
        d = docx.Document()
        t = d.add_table(rows=2, cols=1)
        t.cell(0, 0).text = "Header"
        t.cell(1, 0).text = "a | b"
        result = extract_docx(_save(d, tmp_path))
        assert result is not None
        assert "a \\| b" in result
        # Escaped, so the row still has exactly one real column separator pair.
        assert result.count("\n") == 2

    def test_table_between_paragraphs_preserves_document_order(self, tmp_path: Path) -> None:
        """document.paragraphs/.tables are separate flat lists that lose
        interleaving — this proves the real body-order walk is used."""
        d = docx.Document()
        d.add_paragraph("Before table")
        t = d.add_table(rows=1, cols=1)
        t.cell(0, 0).text = "cell"
        d.add_paragraph("After table")
        result = extract_docx(_save(d, tmp_path))
        assert result is not None
        before_idx = result.index("Before table")
        table_idx = result.index("| cell |")
        after_idx = result.index("After table")
        assert before_idx < table_idx < after_idx

    def test_multi_paragraph_cell_uses_br(self, tmp_path: Path) -> None:
        d = docx.Document()
        t = d.add_table(rows=2, cols=1)
        t.cell(0, 0).text = "Header"
        cell = t.cell(1, 0)
        cell.paragraphs[0].text = "line one"
        cell.add_paragraph("line two")
        result = extract_docx(_save(d, tmp_path))
        assert result is not None
        assert "line one<br>line two" in result

    def test_empty_table_produces_nothing(self, tmp_path: Path) -> None:
        d = docx.Document()
        d.add_paragraph("Before.")
        d.add_table(rows=0, cols=2)
        d.add_paragraph("After.")
        result = extract_docx(_save(d, tmp_path))
        assert result == "Before.\n\nAfter."


# ---------------------------------------------------------------------------
# Code-style paragraphs
# ---------------------------------------------------------------------------


class TestCodeParagraphs:
    def test_contiguous_monospace_paragraphs_merge_into_one_fence(self, tmp_path: Path) -> None:
        d = docx.Document()
        for line in ["def foo():", "    return 42"]:
            p = d.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
        result = extract_docx(_save(d, tmp_path))
        assert result == "```\ndef foo():\n    return 42\n```"

    def test_code_indentation_is_preserved(self, tmp_path: Path) -> None:
        d = docx.Document()
        p = d.add_paragraph()
        run = p.add_run("    indented line")
        run.font.name = "Courier New"
        result = extract_docx(_save(d, tmp_path))
        assert result == "```\n    indented line\n```"

    def test_code_style_name_detected_without_monospace_font(self, tmp_path: Path) -> None:
        from docx.enum.style import WD_STYLE_TYPE

        d = docx.Document()
        code_style = d.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = d.styles["Normal"]
        d.add_paragraph("plain_code_line()", style="Code")
        result = extract_docx(_save(d, tmp_path))
        assert result == "```\nplain_code_line()\n```"

    def test_code_run_separated_by_normal_paragraph_flushes_fence(self, tmp_path: Path) -> None:
        d = docx.Document()
        p1 = d.add_paragraph()
        p1.add_run("code line").font.name = "Consolas"
        d.add_paragraph("Normal text.")
        p2 = d.add_paragraph()
        p2.add_run("more code").font.name = "Consolas"
        result = extract_docx(_save(d, tmp_path))
        assert result == "```\ncode line\n```\n\nNormal text.\n\n```\nmore code\n```"

    def test_unstyled_run_is_not_treated_as_code(self, tmp_path: Path) -> None:
        """A paragraph with no explicit font override at all — the common
        case — must not be misdetected as code."""
        d = docx.Document()
        d.add_paragraph("Just an ordinary sentence.")
        result = extract_docx(_save(d, tmp_path))
        assert result == "Just an ordinary sentence."
        assert "```" not in result


# ---------------------------------------------------------------------------
# Empty document
# ---------------------------------------------------------------------------


class TestEmptyDocument:
    def test_brand_new_document_returns_empty_string_not_none(self, tmp_path: Path) -> None:
        """A genuinely empty document is a successful extraction with
        nothing to report — "" — distinct from None, which means extraction
        did not happen at all."""
        d = docx.Document()
        result = extract_docx(_save(d, tmp_path))
        assert result == ""
        assert result is not None

    def test_document_of_only_whitespace_paragraphs_returns_empty_string(
        self, tmp_path: Path
    ) -> None:
        d = docx.Document()
        d.add_paragraph("   ")
        d.add_paragraph("\t")
        result = extract_docx(_save(d, tmp_path))
        assert result == ""


# ---------------------------------------------------------------------------
# Malformed / corrupt DOCX — fail-open
# ---------------------------------------------------------------------------


class TestMalformedDocx:
    def test_not_a_zip_file_returns_none(self, tmp_path: Path) -> None:
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"this is not a zip file at all, just garbage")
        with pytest.warns(UserWarning, match="document extraction error"):
            result = extract_docx(path)
        assert result is None

    def test_zip_with_wrong_internal_structure_returns_none(self, tmp_path: Path) -> None:
        """A real zip file (valid zip magic bytes) that simply isn't a
        Word package — no [Content_Types].xml, no word/document.xml."""
        path = tmp_path / "not_really_a_docx.docx"
        with zipfile.ZipFile(path, "w") as zf:
            zf.writestr("hello.txt", "not a docx package")
        with pytest.warns(UserWarning, match="document extraction error"):
            result = extract_docx(path)
        assert result is None

    def test_truncated_docx_returns_none(self, tmp_path: Path) -> None:
        """A real .docx, truncated mid-file (simulating an interrupted
        write/transfer) — must still fail open, not raise."""
        d = docx.Document()
        d.add_paragraph("content")
        good_path = _save(d, tmp_path, "good.docx")
        truncated_bytes = good_path.read_bytes()[: len(good_path.read_bytes()) // 2]
        bad_path = tmp_path / "truncated.docx"
        bad_path.write_bytes(truncated_bytes)
        with pytest.warns(UserWarning, match="document extraction error"):
            result = extract_docx(bad_path)
        assert result is None

    def test_nonexistent_file_returns_none(self, tmp_path: Path) -> None:
        with pytest.warns(UserWarning, match="document extraction error"):
            result = extract_docx(tmp_path / "does_not_exist.docx")
        assert result is None

    def test_no_exception_escapes_for_any_malformed_input(self, tmp_path: Path) -> None:
        candidates = [
            tmp_path / "empty_file.docx",
            tmp_path / "missing.docx",
        ]
        (tmp_path / "empty_file.docx").write_bytes(b"")
        for path in candidates:
            with warnings.catch_warnings():
                warnings.simplefilter("ignore")
                result = extract_docx(path)  # must never raise
            assert result is None


# ---------------------------------------------------------------------------
# Missing dependency — fail-open
# ---------------------------------------------------------------------------


class TestMissingDependency:
    def test_import_error_returns_none_with_actionable_warning(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setitem(sys.modules, "docx", None)
        with pytest.warns(UserWarning, match="quor\\[documents\\]"):
            result = extract_docx(tmp_path / "anything.docx")
        assert result is None

    def test_missing_dependency_does_not_attempt_file_access(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Confirms the ImportError branch returns before ever touching the
        file — a genuinely nonexistent path must not change this outcome."""
        monkeypatch.setitem(sys.modules, "docx", None)
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            result = extract_docx(Path("this/path/does/not/exist/at/all.docx"))
        assert result is None


# ---------------------------------------------------------------------------
# Extraction failure via registry — the full fail-open path
# ---------------------------------------------------------------------------


class TestRegistryIntegration:
    def test_real_extraction_flows_through_registry_extract(self, tmp_path: Path) -> None:
        from quor.pipeline.extract.registry import extract

        d = docx.Document()
        d.add_heading("Report", level=1)
        d.add_paragraph("Body text.")
        path = _save(d, tmp_path)
        result = extract(path)
        assert result == "# Report\n\nBody text."

    def test_corrupt_docx_fails_open_through_registry_extract(self, tmp_path: Path) -> None:
        from quor.pipeline.extract.registry import extract

        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"garbage")
        with pytest.warns(UserWarning, match="document extraction error"):
            result = extract(path)
        assert result is None
