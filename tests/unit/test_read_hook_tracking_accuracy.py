"""QB-094: Read-hook tracking-accuracy coverage.

Before this ticket, `track_invocation()` ran inside `_compress_read_output()`
(or one of its two helpers) — before `_handle_text()` had a chance to
prepend the "Repository Context" block (QB-079), the "Relevant repository
files" block (QB-081), the "Repository Tip" nudge (QB-090), or the
concise-instruction prefix. The tracked `final_tokens` therefore always
excluded whatever those layers added, understating (sometimes to zero, or
even the wrong sign) what Claude actually received.

QB-094's fix (Option C) makes producers return a `_ReadCompressionResult`
carrying tracking metadata instead of calling `track_invocation()`
themselves; `_handle_text()` tracks exactly once, after every layer above
has already been applied. Every test below drives the real
`run_hook(tracking=...)` entry point with a fake `TrackingDB` and asserts
the one recorded row's `final_tokens` equals `count_tokens()` of the actual
`updatedToolOutput` bytes — the exact invariant that was violated before
this fix, and the gap that let it ship unnoticed in the first place (no
existing suite asserted on tracked token counts at all).
"""

from __future__ import annotations

import io
import subprocess
import sys
from pathlib import Path
from unittest.mock import patch

import docx
import orjson
import pytest
from reportlab.pdfgen import canvas

from quor.adapters.claude_read import run_hook
from quor.adapters.dispatcher import CONCISE_INSTRUCTION
from quor.pipeline.repo_profile import intel_store
from quor.pipeline.repo_profile.intel_model import FileIntelligenceEntry
from quor.tracking.db import InvocationRecord, count_tokens

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class _FakeStdout:
    def __init__(self) -> None:
        self.buffer: io.BytesIO = io.BytesIO()

    def write(self, s: str) -> int:
        return 0

    def flush(self) -> None:
        pass


class _FakeTracking:
    """Stand-in for `TrackingDB`: `track_invocation()` only ever calls
    `.record(rec)` on whatever it's given (see `quor/tracking/db.py`), so a
    plain list-appending fake is enough to inspect exactly what would have
    been persisted, with no real SQLite involved."""

    def __init__(self) -> None:
        self.records: list[InvocationRecord] = []

    def record(self, rec: InvocationRecord) -> None:
        self.records.append(rec)


def _read_payload(file_path: str, tool_response: str, *, transcript_path: str = "") -> dict:
    payload: dict = {
        "tool_name": "Read",
        "tool_input": {"file_path": file_path},
        "tool_response": tool_response,
    }
    if transcript_path:
        payload["transcript_path"] = transcript_path
    return payload


def _run_hook_tracked(payload: dict) -> tuple[dict, _FakeTracking]:
    """Drive the real `run_hook()` entry point with a fake `TrackingDB` and
    return both the parsed hook response and the fake's recorded rows."""
    raw = orjson.dumps(payload).decode("utf-8")
    fake_stdout = _FakeStdout()
    tracking = _FakeTracking()
    with (
        patch.object(sys, "stdin", io.StringIO(raw)),
        patch.object(sys, "stdout", fake_stdout),
    ):
        run_hook(tracking=tracking)
    fake_stdout.buffer.seek(0)
    response = orjson.loads(fake_stdout.buffer.read())
    return response, tracking


def _write_transcript(tmp_path: Path, *user_messages: str, name: str = "transcript.jsonl") -> Path:
    transcript = tmp_path / name
    lines = [
        orjson.dumps(
            {
                "type": "user",
                "message": {"role": "user", "content": [{"type": "text", "text": text}]},
            }
        ).decode()
        for text in user_messages
    ]
    transcript.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return transcript


def _write_intel_entry(root: Path, rel_path: str, **overrides: object) -> None:
    fields: dict[str, object] = {
        "language": "python",
        "kind": "source",
        "importance": "Low",
        "top_symbols": [],
    }
    fields.update(overrides)
    existing = intel_store.load_file_intelligence(root) or {}
    existing[rel_path] = FileIntelligenceEntry(**fields)  # type: ignore[arg-type]
    intel_store.save_file_intelligence(root, existing)


def _init_git_repo(root: Path) -> None:
    subprocess.run(["git", "init", "-q"], cwd=root, check=True)
    subprocess.run(["git", "config", "user.email", "test@example.com"], cwd=root, check=True)
    subprocess.run(["git", "config", "user.name", "Test"], cwd=root, check=True)
    subprocess.run(["git", "add", "."], cwd=root, check=True)
    subprocess.run(["git", "commit", "-q", "-m", "init"], cwd=root, check=True)


_PYTHON_SOURCE = '''import os

DEFAULT_TIMEOUT = 30


def fetch_data(url, timeout=DEFAULT_TIMEOUT):
    """Fetch data from a URL."""
    response = make_request(url, timeout=timeout)
    if response.status_code != 200:
        raise RuntimeError("bad response")
    return response.json()


class Client:
    def __init__(self, base_url):
        self.base_url = base_url

    def get(self, path):
        full_url = self.base_url + path
        return fetch_data(full_url)
'''


def _write_matching_source(root: Path, rel_path: str, **overrides: object) -> Path:
    """Write a real `.py` file plus a `FileIntelligenceEntry` whose
    size/mtime_ns match it (QB-079's staleness check), mirroring
    tests/unit/test_read_hook_repo_context.py's own fixture exactly."""
    path = root / rel_path
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(_PYTHON_SOURCE, encoding="utf-8")
    st = path.stat()
    fields: dict[str, object] = {
        "language": "python",
        "kind": "source",
        "importance": "High",
        "imports": 3,
        "imported_by": 61,
        "entry_point": False,
        "top_symbols": ["fetch_data", "Client"],
        "size": st.st_size,
        "mtime_ns": st.st_mtime_ns,
    }
    fields.update(overrides)
    existing = intel_store.load_file_intelligence(root) or {}
    existing[rel_path] = FileIntelligenceEntry(**fields)  # type: ignore[arg-type]
    intel_store.save_file_intelligence(root, existing)
    return path


def _large_docx(tmp_path: Path, name: str = "report.docx") -> Path:
    d = docx.Document()
    d.add_heading("Design Notes", level=1)
    d.add_paragraph("REQ-1: must survive extraction and compression.")
    for _ in range(150):
        d.add_paragraph("This is an ordinary sentence of filler prose repeated many times.")
    path = tmp_path / name
    d.save(str(path))
    return path


def _large_pdf(tmp_path: Path, name: str = "report.pdf") -> Path:
    path = tmp_path / name
    c = canvas.Canvas(str(path), pagesize=(500, 3000))
    c.setFont("Helvetica-Bold", 20)
    c.drawString(72, 2950, "Design Notes")
    c.setFont("Helvetica", 11)
    y = 2910
    c.drawString(72, y, "REQ-1: must survive extraction and compression.")
    y -= 20
    for _ in range(150):
        c.drawString(72, y, "This is an ordinary sentence of filler prose repeated many times.")
        y -= 16
    c.save()
    return path


def _assert_tracking_matches_output(response: dict, tracking: _FakeTracking) -> InvocationRecord:
    """The core QB-094 invariant: exactly one row recorded, and its
    `final_tokens` equals `count_tokens()` of whatever `updatedToolOutput`
    (or, if omitted, nothing changed) actually is."""
    assert len(tracking.records) == 1
    rec = tracking.records[0]
    updated = response["hookSpecificOutput"].get("updatedToolOutput")
    expected_final = count_tokens(updated) if isinstance(updated, str) else rec.original_tokens
    assert rec.final_tokens == expected_final
    return rec


# ---------------------------------------------------------------------------
# 1. Concise instruction only
# ---------------------------------------------------------------------------


class TestConciseInstructionOnly:
    def test_tracked_tokens_include_the_instruction(self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.chdir(tmp_path)
        large_markdown = "# Heading\n\n" + ("This is filler prose. " * 2000)
        response, tracking = _run_hook_tracked(_read_payload("notes.md", large_markdown))

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert updated.startswith(CONCISE_INSTRUCTION)
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.filter_name == "markdown"
        assert rec.was_passthrough is False


# ---------------------------------------------------------------------------
# 2. Repository Context only (QB-079)
# ---------------------------------------------------------------------------


class TestRepositoryContextOnly:
    def test_tracked_tokens_include_the_context_block(self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.chdir(tmp_path)
        path = _write_matching_source(tmp_path, "app.py")
        # No transcript_path -> relevant-files and the repo-intel nudge both
        # stay inert (both require it), isolating Repository Context alone.
        response, tracking = _run_hook_tracked(_read_payload(str(path), _PYTHON_SOURCE))

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert "Repository Context (app.py)" in updated
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.filter_name == "cat-python"


# ---------------------------------------------------------------------------
# 3. Relevant Repository Files only (QB-081)
# ---------------------------------------------------------------------------


class TestRelevantFilesOnly:
    def test_tracked_tokens_include_the_relevant_files_block(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.chdir(tmp_path)
        # .rs has no compression filter at all, and no git repo exists here
        # (compute_hook_nudge requires one), so this isolates QB-081 alone.
        unsupported_source = 'fn main() {\n    println!("hello");\n}\n'
        (tmp_path / "main.rs").write_text(unsupported_source, encoding="utf-8")
        _write_intel_entry(tmp_path, "src/auth/login.py", top_symbols=["LoginManager"])
        transcript = _write_transcript(tmp_path, "Where is LoginManager defined?")

        response, tracking = _run_hook_tracked(
            _read_payload(str(tmp_path / "main.rs"), unsupported_source, transcript_path=str(transcript))
        )

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert "Relevant repository files" in updated
        rec = _assert_tracking_matches_output(response, tracking)
        # No filter matched — this is the exact "passthrough that still grew"
        # case QB-094 exists to fix; see TestPassthroughWithEnhancement below
        # for the was_passthrough/negative-delta assertions specifically.
        assert rec.filter_name is None
        assert rec.was_passthrough is True


# ---------------------------------------------------------------------------
# 4. Repository Tip nudge only (QB-090)
# ---------------------------------------------------------------------------


class TestRepositoryTipOnly:
    def test_tracked_tokens_include_the_nudge(self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
        root = tmp_path / "repo"
        root.mkdir()
        (root / "unrelated.xyz").write_text("nothing quor filters\n", encoding="utf-8")
        _init_git_repo(root)
        monkeypatch.chdir(root)

        # transcript_path points at a file that doesn't exist, so relevant-
        # files finds no prompt text and stays inert; no file_intelligence.json
        # was ever written, so repo context can't apply either — isolating
        # the nudge alone, same setup as test_read_hook_repo_intel_nudge.py.
        response, tracking = _run_hook_tracked(
            _read_payload(str(root / "unrelated.xyz"), "nothing quor filters\n", transcript_path="transcript.jsonl")
        )

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert "Repository Tip" in updated
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.filter_name is None
        assert rec.was_passthrough is True


# ---------------------------------------------------------------------------
# 5. Combined — maximum layering
# ---------------------------------------------------------------------------


class TestCombinedLayering:
    def test_tracked_tokens_match_output_with_every_layer_active(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        root = tmp_path / "repo"
        root.mkdir()
        monkeypatch.chdir(root)
        path = _write_matching_source(root, "app.py")
        _write_intel_entry(root, "src/auth/login.py", top_symbols=["LoginManager"])
        transcript = _write_transcript(root, "Where is LoginManager defined?")
        _init_git_repo(root)  # enables the repo-intel nudge too

        response, tracking = _run_hook_tracked(
            _read_payload(str(path), _PYTHON_SOURCE, transcript_path=str(transcript))
        )

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert isinstance(updated, str)
        # At least Repository Context (QB-079, applied inside the producer)
        # and Relevant repository files (QB-081, applied in _handle_text)
        # are both genuinely present — proving the tracked total spans a
        # producer-level prepend and a _handle_text-level prepend at once.
        assert "Repository Context (app.py)" in updated
        assert "Relevant repository files" in updated
        _assert_tracking_matches_output(response, tracking)


# ---------------------------------------------------------------------------
# 6. Passthrough with enhancement (previously recorded a false zero delta)
# ---------------------------------------------------------------------------


class TestPassthroughWithEnhancement:
    def test_was_passthrough_stays_true_while_tokens_reflect_the_addition(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.chdir(tmp_path)
        unsupported_source = 'fn main() {\n    println!("hello");\n}\n'
        (tmp_path / "main.rs").write_text(unsupported_source, encoding="utf-8")
        _write_intel_entry(tmp_path, "src/auth/login.py", top_symbols=["LoginManager"])
        transcript = _write_transcript(tmp_path, "Where is LoginManager defined?")

        response, tracking = _run_hook_tracked(
            _read_payload(str(tmp_path / "main.rs"), unsupported_source, transcript_path=str(transcript))
        )

        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.was_passthrough is True
        assert rec.filter_name is None
        assert rec.original_tokens == count_tokens(unsupported_source)
        # The bug this ticket fixes: before QB-094 this row recorded
        # final_tokens == original_tokens (a false zero delta), even though
        # Claude genuinely received the original text plus a whole
        # "Relevant repository files" block on top.
        assert rec.final_tokens > rec.original_tokens


# ---------------------------------------------------------------------------
# 7. Pure passthrough — still tracked exactly once
# ---------------------------------------------------------------------------


class TestPurePassthrough:
    def test_tracked_once_with_no_change(self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
        monkeypatch.chdir(tmp_path)
        content = 'fn main() {\n    println!("hello");\n}\n'
        response, tracking = _run_hook_tracked(_read_payload("main.rs", content))

        assert "updatedToolOutput" not in response["hookSpecificOutput"]
        assert len(tracking.records) == 1
        rec = tracking.records[0]
        assert rec.original_tokens == rec.final_tokens == count_tokens(content)
        assert rec.was_passthrough is True
        assert rec.filter_name is None


# ---------------------------------------------------------------------------
# 8. Extraction success
# ---------------------------------------------------------------------------


class TestExtractionSuccess:
    def test_docx_tracked_tokens_match_extracted_and_compressed_output(self, tmp_path: Path) -> None:
        path = _large_docx(tmp_path)
        response, tracking = _run_hook_tracked(_read_payload(str(path), "<binary Read result placeholder>"))

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert "REQ-1: must survive extraction and compression." in updated
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.filter_name == "markdown"
        assert rec.original_tokens == count_tokens("<binary Read result placeholder>")

    def test_pdf_tracked_tokens_match_extracted_and_compressed_output(self, tmp_path: Path) -> None:
        path = _large_pdf(tmp_path)
        response, tracking = _run_hook_tracked(_read_payload(str(path), "<binary Read result placeholder>"))

        updated = response["hookSpecificOutput"]["updatedToolOutput"]
        assert "REQ-1: must survive extraction and compression." in updated
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.filter_name == "markdown"


# ---------------------------------------------------------------------------
# 9. Extraction failure
# ---------------------------------------------------------------------------


class TestExtractionFailure:
    def test_corrupt_docx_tracks_original_unchanged(self, tmp_path: Path) -> None:
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"not a real docx file")
        with pytest.warns(UserWarning, match="document extraction error"):
            response, tracking = _run_hook_tracked(_read_payload(str(path), "original content"))

        assert "updatedToolOutput" not in response["hookSpecificOutput"]
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.was_passthrough is True
        assert rec.filter_name is None
        assert rec.original_tokens == rec.final_tokens == count_tokens("original content")

    def test_nonexistent_docx_path_tracks_original_unchanged(self, tmp_path: Path) -> None:
        missing = tmp_path / "does_not_exist.docx"
        response, tracking = _run_hook_tracked(_read_payload(str(missing), "original content"))

        assert "updatedToolOutput" not in response["hookSpecificOutput"]
        rec = _assert_tracking_matches_output(response, tracking)
        assert rec.was_passthrough is True
        assert rec.original_tokens == rec.final_tokens == count_tokens("original content")
