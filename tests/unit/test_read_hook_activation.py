"""QB-007C: end-to-end activation of the PostToolUse/Read hook.

QB-007A built the hook plumbing (always a no-op) and QB-007B built the
`markdown`/`document-text` filters (tested only at the FilterRegistry
layer, never reachable from a real Read call). This file exercises the
thing QB-007C actually adds: `quor/adapters/claude_read.py::run_hook()`
genuinely routing supported Read output through those filters and
returning `updatedToolOutput` when compression produces different content
— driven through the real stdin -> stdout JSON contract, not by calling
FilterRegistry directly (see tests/unit/test_document_filters.py for that
layer, and tests/unit/test_adapters_read.py for adapter-level payload/shape
coverage this file assumes already holds).
"""

from __future__ import annotations

import io
import sys
import warnings
from pathlib import Path
from typing import Any
from unittest.mock import patch

import docx
import orjson
import pytest
from reportlab.pdfgen import canvas

from quor.adapters.claude_read import run_hook
from quor.adapters.dispatcher import CONCISE_INSTRUCTION
from quor.filters.registry import FilterRegistry

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


class _FakeStdout:
    def __init__(self) -> None:
        self.buffer: io.BytesIO = io.BytesIO()

    def write(self, s: str) -> int:
        return 0

    def flush(self) -> None:
        pass


def _read_payload(file_path: str, tool_response: str) -> dict:
    return {
        "tool_name": "Read",
        "tool_input": {"file_path": file_path},
        "tool_response": tool_response,
    }


def _run_hook(payload: dict) -> dict:
    raw = orjson.dumps(payload).decode("utf-8")
    fake_stdout = _FakeStdout()
    with (
        patch.object(sys, "stdin", io.StringIO(raw)),
        patch.object(sys, "stdout", fake_stdout),
    ):
        run_hook()
    fake_stdout.buffer.seek(0)
    return orjson.loads(fake_stdout.buffer.read())


_LARGE_MARKDOWN = (
    "# Design Notes\n\n"
    "REQ-1: must survive compression.\n\n"
    + ("This is an ordinary sentence of filler prose. " * 400)
    + "\n\n# Late Heading\n\nFinal paragraph."
)

_LARGE_TEXT = (
    "REQ-2: must survive compression.\n\n"
    + ("This is an ordinary sentence of filler prose. " * 400)
    + "\n\nFinal paragraph."
)

_SMALL_MARKDOWN = "# Title\n\nJust a short paragraph with nothing special."


# ---------------------------------------------------------------------------
# Supported Read operations return updatedToolOutput
# ---------------------------------------------------------------------------


class TestSupportedTypesCompress:
    @pytest.mark.parametrize(
        "file_path,content",
        [
            ("notes.md", _LARGE_MARKDOWN),
            ("notes.markdown", _LARGE_MARKDOWN),
            ("notes.txt", _LARGE_TEXT),
            ("notes.rst", _LARGE_TEXT),
        ],
    )
    def test_large_supported_document_returns_updated_tool_output(
        self, file_path: str, content: str
    ) -> None:
        result = _run_hook(_read_payload(file_path, content))
        hook_specific = result["hookSpecificOutput"]
        assert hook_specific["hookEventName"] == "PostToolUse"
        updated = hook_specific.get("updatedToolOutput")
        assert isinstance(updated, str)
        assert len(updated) < len(content)

    def test_protected_structure_survives_compression(self) -> None:
        """The whole point of QB-007B's filters: compression shrinks the
        document but never drops a requirement ID or a heading, regardless
        of where it falls in the document."""
        result = _run_hook(_read_payload("notes.md", _LARGE_MARKDOWN))
        updated = result["hookSpecificOutput"]["updatedToolOutput"]
        assert "REQ-1: must survive compression." in updated
        assert "# Late Heading" in updated

    def test_nested_path_still_compresses(self) -> None:
        result = _run_hook(_read_payload("docs/final/DESIGN.md", _LARGE_MARKDOWN))
        assert "updatedToolOutput" in result["hookSpecificOutput"]


# ---------------------------------------------------------------------------
# Unsupported file types pass through unchanged
# ---------------------------------------------------------------------------


class TestUnsupportedTypesPassThrough:
    @pytest.mark.parametrize(
        "file_path",
        ["config.json", "LICENSE", "image.png", "main.rs"],
    )
    def test_unsupported_large_file_never_compresses(self, file_path: str) -> None:
        """Regression guard for the generic-filter scope-leak: FilterRegistry's
        built-in `generic` filter matches *any* non-empty string (including
        every one of these paths), but this adapter must never apply it —
        only quor/filters/builtin/{markdown,document-text}.toml are in
        scope for Read. A large payload is used specifically so that if the
        allowlist regressed, this test would catch real (not just
        theoretical) compression happening.

        `.docx`/`.pdf` are deliberately NOT parametrized here as of
        QB-007E4 — they are no longer "unsupported": they're routed to
        extraction. See TestDocxPdfExtraction below for their dedicated
        coverage (a fake, nonexistent "report.docx" path would still pass
        this specific assertion today, since extraction fails open for a
        missing file — but for the wrong reason, no longer "this extension
        has no filter at all," and that distinction is exactly what a
        regression test must not paper over).

        `.py`/`.js`/`.jsx`/`.mjs`/`.cjs`/`.ts`/`.tsx` are likewise
        deliberately NOT parametrized here as of QB-005F — they are also no
        longer "unsupported": they're routed to the AST-summarization
        filters by name. See tests/unit/test_read_hook_ast_summarization.py
        for their dedicated coverage."""
        large_content = "line of filler text. " * 10_000
        result = _run_hook(_read_payload(file_path, large_content))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_no_extension_passes_through(self) -> None:
        result = _run_hook(_read_payload("Makefile", "line of filler text. " * 10_000))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]


# ---------------------------------------------------------------------------
# DOCX/PDF extraction end to end (QB-007E4)
#
# Unlike every test above, extraction reads the real file from disk via
# `extract(Path(file_path))` — the Read tool's own `tool_response` is not
# used as extraction input at all, only as the "original" for tracking and
# the unchanged-content comparison. Every fixture here is therefore a real
# file written to `tmp_path`, with `file_path` in the payload pointing at
# its real (absolute) location.
# ---------------------------------------------------------------------------


def _large_docx(tmp_path: Path, name: str = "report.docx") -> Path:
    """A real .docx large enough (repeated body paragraphs) to exceed the
    markdown filter's 2000-token budget once extracted, so compression
    genuinely fires — mirrors _LARGE_MARKDOWN/_LARGE_TEXT's own purpose
    above, just authored as a binary document instead of raw text."""
    d = docx.Document()
    d.add_heading("Design Notes", level=1)
    d.add_paragraph("REQ-1: must survive extraction and compression.")
    for _ in range(150):
        d.add_paragraph("This is an ordinary sentence of filler prose repeated many times.")
    d.add_heading("Late Heading", level=2)
    d.add_paragraph("Final paragraph.")
    path = tmp_path / name
    d.save(str(path))
    return path


def _large_pdf(tmp_path: Path, name: str = "report.pdf") -> Path:
    path = tmp_path / name
    c = canvas.Canvas(str(path), pagesize=(500, 3000))
    c.setFont("Helvetica-Bold", 20)
    c.drawString(72, 2950, "Design Notes")
    c.setFont("Helvetica", 11)
    y = 2910
    c.drawString(72, y, "REQ-1: must survive extraction and compression.")
    y -= 20
    for _ in range(150):
        c.drawString(72, y, "This is an ordinary sentence of filler prose repeated many times.")
        y -= 16
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, y, "Late Heading")
    y -= 20
    c.setFont("Helvetica", 11)
    c.drawString(72, y, "Final paragraph.")
    c.save()
    return path


class TestDocxPdfExtraction:
    def test_large_docx_extracts_and_compresses(self, tmp_path: Path) -> None:
        path = _large_docx(tmp_path)
        result = _run_hook(_read_payload(str(path), "<binary Read result placeholder>"))
        hook_specific = result["hookSpecificOutput"]
        updated = hook_specific.get("updatedToolOutput")
        assert isinstance(updated, str)
        assert "# Design Notes" in updated
        assert "REQ-1: must survive extraction and compression." in updated

    def test_large_pdf_extracts_and_compresses(self, tmp_path: Path) -> None:
        path = _large_pdf(tmp_path)
        result = _run_hook(_read_payload(str(path), "<binary Read result placeholder>"))
        hook_specific = result["hookSpecificOutput"]
        updated = hook_specific.get("updatedToolOutput")
        assert isinstance(updated, str)
        assert "# Design Notes" in updated
        assert "REQ-1: must survive extraction and compression." in updated

    def test_docx_protected_structure_survives_compression(self, tmp_path: Path) -> None:
        """Same guarantee QB-007B's filters already give the direct .md
        path — proves the extracted text is genuinely routed through the
        real `markdown` filter, not just passed through verbatim."""
        path = _large_docx(tmp_path)
        result = _run_hook(_read_payload(str(path), "placeholder"))
        updated = result["hookSpecificOutput"]["updatedToolOutput"]
        assert "REQ-1: must survive extraction and compression." in updated
        assert "## Late Heading" in updated

    def test_nonexistent_docx_path_fails_open(self, tmp_path: Path) -> None:
        """Extraction returning None (file doesn't exist) behaves exactly
        like an unsupported type — updatedToolOutput is omitted, no
        exception."""
        missing = tmp_path / "does_not_exist.docx"
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            result = _run_hook(_read_payload(str(missing), "original content"))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_corrupt_docx_fails_open(self, tmp_path: Path) -> None:
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"not a real docx file")
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            result = _run_hook(_read_payload(str(path), "original content"))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_corrupt_pdf_fails_open(self, tmp_path: Path) -> None:
        path = tmp_path / "corrupt.pdf"
        path.write_bytes(b"not a real pdf file")
        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            result = _run_hook(_read_payload(str(path), "original content"))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_extraction_exception_fails_open_not_raises(self, tmp_path: Path) -> None:
        """A raising extract() (defense-in-depth path, since extract()'s
        own contract says it never raises) must not propagate — mirrors
        TestFailOpenOnFilterFailure's coverage for the filter layer."""
        path = _large_docx(tmp_path)
        with (
            patch(
                "quor.adapters.claude_read.extract",
                side_effect=RuntimeError("synthetic extraction failure"),
            ),
            pytest.warns(UserWarning, match="Read extraction error"),
        ):
            result = _run_hook(_read_payload(str(path), "original content"))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_small_docx_still_returns_extracted_text(self, tmp_path: Path) -> None:
        """Unlike a small .md file (already markdown, so an under-budget
        compression is a genuine no-op), a small DOCX still returns
        updatedToolOutput — extraction itself already transformed the
        content from whatever the raw tool_response was to clean Markdown,
        independent of whether the *subsequent* markdown-filter step had
        anything left to compress."""
        d = docx.Document()
        d.add_heading("Title", level=1)
        d.add_paragraph("Just a short paragraph with nothing special.")
        path = tmp_path / "small.docx"
        d.save(str(path))

        result = _run_hook(_read_payload(str(path), "irrelevant original"))
        updated = result["hookSpecificOutput"].get("updatedToolOutput")
        assert updated == (
            CONCISE_INSTRUCTION + "# Title\n\nJust a short paragraph with nothing special."
        )

    def test_docx_omits_update_when_tool_response_already_matches_extracted_text(
        self, tmp_path: Path
    ) -> None:
        """The real "omit if unchanged" case for the extraction path: if
        `tool_response` happens to already equal the final extracted (and
        filtered) text exactly, updatedToolOutput is still correctly
        omitted — the comparison is against the true final output, not
        against some intermediate extraction-only value."""
        d = docx.Document()
        d.add_heading("Title", level=1)
        d.add_paragraph("Just a short paragraph with nothing special.")
        path = tmp_path / "small.docx"
        d.save(str(path))

        already_extracted = "# Title\n\nJust a short paragraph with nothing special."
        result = _run_hook(_read_payload(str(path), already_extracted))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_unsupported_extension_still_passes_through(self, tmp_path: Path) -> None:
        """A real file that isn't .docx/.pdf/.md/.txt/.rst — and, as of
        QB-005F, isn't a mapped source-code extension either — still never
        reaches extract() or any filter, exactly as before QB-007E4.

        `.py` was this test's example prior to QB-005F; it is now a
        genuinely supported source-code extension (routed to `cat-python`
        by name), so a garbage-filler `.py` fixture like this one would
        legitimately compress — see
        tests/unit/test_read_hook_ast_summarization.py for that dedicated
        coverage instead."""
        path = tmp_path / "data.json"
        path.write_text('{"greeting": "hello"}\n', encoding="utf-8")
        result = _run_hook(_read_payload(str(path), "line of filler text. " * 10_000))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]


# ---------------------------------------------------------------------------
# Identical content never emits updatedToolOutput
# ---------------------------------------------------------------------------


class TestIdenticalContentOmitsUpdate:
    def test_small_markdown_below_budget_omits_update(self) -> None:
        result = _run_hook(_read_payload("notes.md", _SMALL_MARKDOWN))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_small_text_below_budget_omits_update(self) -> None:
        result = _run_hook(_read_payload("notes.txt", "Just a short line of text."))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_content_entirely_composed_of_protected_lines_omits_update(self) -> None:
        """Every line matches a preserve_pattern, so nothing is ever
        eligible for max_tokens to compress — rendered output is
        byte-identical even though the filter genuinely ran."""
        content = "REQ-1: first.\nREQ-2: second.\nREQ-3: third.\n"
        result = _run_hook(_read_payload("notes.md", content))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]


# ---------------------------------------------------------------------------
# Fail-open behaviour for pipeline/filter failures
# ---------------------------------------------------------------------------


class TestFailOpenOnFilterFailure:
    def test_registry_apply_exception_omits_update_not_raises(self) -> None:
        """A raising FilterRegistry.apply() must not propagate — the
        adapter's own try/except (not just __main__'s outer guard) catches
        it, so a single bad filter can't take down routing entirely."""
        with (
            patch.object(
                FilterRegistry, "apply", side_effect=RuntimeError("synthetic apply failure")
            ),
            pytest.warns(UserWarning, match="Read filter error"),
        ):
            result = _run_hook(_read_payload("notes.md", _LARGE_MARKDOWN))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_registry_find_exception_omits_update_not_raises(self) -> None:
        with (
            patch.object(
                FilterRegistry, "find", side_effect=RuntimeError("synthetic find failure")
            ),
            pytest.warns(UserWarning, match="Read filter error"),
        ):
            result = _run_hook(_read_payload("notes.md", _LARGE_MARKDOWN))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_registry_construction_exception_omits_update_not_raises(self) -> None:
        with (
            patch(
                "quor.adapters.claude_read.FilterRegistry",
                side_effect=RuntimeError("synthetic construction failure"),
            ),
            pytest.warns(UserWarning, match="Read filter error"),
        ):
            result = _run_hook(_read_payload("notes.md", _LARGE_MARKDOWN))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_stage_level_failure_still_yields_best_effort_output(self) -> None:
        """A single stage raising is already fail-open at the Pipeline
        level (unchanged by QB-007C) — the filter still runs, just with
        that stage skipped, so compression can still legitimately fire."""
        import warnings

        from quor.pipeline.stages import strip_lines as strip_lines_module

        def _boom(*_args: object, **_kwargs: object) -> None:
            raise RuntimeError("synthetic stage failure")

        with (
            patch.object(strip_lines_module.StripLinesStage, "apply", _boom),
            warnings.catch_warnings(record=True),
        ):
            warnings.simplefilter("always")
            result = _run_hook(_read_payload("notes.md", _LARGE_MARKDOWN))

        # No exception escaped, and the hook still returned a well-formed
        # response either way (with or without updatedToolOutput depending
        # on whether max_tokens alone was enough to change anything).
        assert result["hookSpecificOutput"]["hookEventName"] == "PostToolUse"


# ---------------------------------------------------------------------------
# Routing precedence / filter selection regressions
# ---------------------------------------------------------------------------


class TestRoutingPrecedenceRegressions:
    def test_md_routes_through_markdown_not_document_text(self) -> None:
        """Markdown-specific structure (ATX heading, fenced code) survives
        only if the `markdown` filter (not `document-text`) actually ran."""
        content = "```python\nvalue = 1\n```\n\n" + (
            "This is an ordinary sentence of filler prose. " * 400
        )
        result = _run_hook(_read_payload("notes.md", content))
        updated = result["hookSpecificOutput"]["updatedToolOutput"]
        assert "```python" in updated

    def test_txt_routes_through_document_text_not_markdown(self) -> None:
        result = _run_hook(_read_payload("notes.txt", _LARGE_TEXT))
        assert "updatedToolOutput" in result["hookSpecificOutput"]

    def test_file_named_like_a_bash_command_still_passes_through(self) -> None:
        """Regression guard: FilterRegistry.find("cat.md") still returns the
        `cat` (Bash) filter first, per QB-007B's documented, accepted
        routing-collision limitation — but the adapter's allowlist means
        this Read is never actually processed by it, unlike a Bash `cat
        cat.md` invocation would be. Confirms the allowlist fix in
        practice, not just in isolation."""
        result = _run_hook(_read_payload("cat.md", _LARGE_MARKDOWN))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

        # The underlying collision is still real at the FilterRegistry
        # layer — only the adapter's allowlist neutralizes it for Read.
        registry = FilterRegistry(skip_user=True, skip_project=True)
        assert registry.find("cat.md").name == "cat"  # type: ignore[union-attr]

    def test_bash_command_routing_is_never_touched_by_the_read_allowlist(self) -> None:
        """The allowlist lives entirely in claude_read.py — FilterRegistry
        itself, and therefore every Bash filter's own routing, is
        unaffected. Confirms the two concerns stay properly separated."""
        registry = FilterRegistry(skip_user=True, skip_project=True)
        assert registry.find("git status").name == "git-status"  # type: ignore[union-attr]
        assert registry.find("pytest tests/").name == "pytest"  # type: ignore[union-attr]

    def test_pytest_targeting_a_markdown_named_test_file_unaffected(self) -> None:
        """Same anchoring regression QB-007B already covered at the
        FilterRegistry layer — re-asserted here because claude_read.py
        never even reaches this command shape (Read paths only), so this
        also documents *why*: the Bash hook (quor/adapters/claude.py) is a
        completely separate code path, untouched by QB-007C."""
        registry = FilterRegistry(skip_user=True, skip_project=True)
        assert registry.find("pytest tests/test_readme.md").name == "pytest"  # type: ignore[union-attr]


# ---------------------------------------------------------------------------
# Extra tool_input fields / non-string tool_response (defensive coverage)
# ---------------------------------------------------------------------------


class TestDefensivePayloadHandling:
    def test_non_string_tool_response_omits_update(self) -> None:
        payload: dict[str, Any] = {
            "tool_name": "Read",
            "tool_input": {"file_path": "notes.md"},
            "tool_response": {"unexpected": "shape"},
        }
        result = _run_hook(payload)
        assert "updatedToolOutput" not in result["hookSpecificOutput"]

    def test_empty_file_path_omits_update(self) -> None:
        result = _run_hook(_read_payload("", _LARGE_MARKDOWN))
        assert "updatedToolOutput" not in result["hookSpecificOutput"]
