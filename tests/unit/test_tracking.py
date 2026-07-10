"""Unit tests for quor/tracking/db.py — SQLite + JSONL persistence."""

from __future__ import annotations

import io
import os
import sqlite3
import subprocess
import threading
import time
import warnings
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import orjson
import pytest

from quor.tracking.db import (
    InvocationRecord,
    TrackingDB,
    count_tokens,
    get_tracking_db,
    normalize_project_path,
    query_gain,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _db_and_jsonl(tmp_path: Path) -> tuple[TrackingDB, Path, Path]:
    db_path = tmp_path / "quor.db"
    jsonl_path = tmp_path / "invocations.jsonl"
    db = TrackingDB(db_path=db_path, jsonl_path=jsonl_path)
    return db, db_path, jsonl_path


def _sample_record(**kwargs) -> InvocationRecord:
    defaults = {
        "command": "git status",
        "project_path": "/home/user/myproject",
        "original_tokens": 100,
        "final_tokens": 20,
        "filter_name": "git-status",
        "was_passthrough": False,
        "duration_ms": 12.5,
    }
    defaults.update(kwargs)
    return InvocationRecord(**defaults)


# ---------------------------------------------------------------------------
# count_tokens
# ---------------------------------------------------------------------------


class TestCountTokens:
    def test_empty_string(self) -> None:
        assert count_tokens("") == 0

    def test_four_chars(self) -> None:
        assert count_tokens("abcd") == 1

    def test_five_chars_rounds_up(self) -> None:
        assert count_tokens("abcde") == 2

    def test_large_text(self) -> None:
        text = "a" * 400
        assert count_tokens(text) == 100

    def test_non_ascii(self) -> None:
        # Unicode chars count by character, not bytes
        assert count_tokens("héllo") == 2  # ceil(5/4)


# ---------------------------------------------------------------------------
# InvocationRecord
# ---------------------------------------------------------------------------


class TestInvocationRecord:
    def test_to_dict_fields(self) -> None:
        rec = _sample_record()
        d = rec.to_dict()
        assert d["command"] == "git status"
        assert d["project_path"] == "/home/user/myproject"
        assert d["original_tokens"] == 100
        assert d["final_tokens"] == 20
        assert d["filter_name"] == "git-status"
        assert d["was_passthrough"] == 0  # bool → int
        assert d["duration_ms"] == 12.5
        assert d["schema_version"] == 2  # bumped for project_key_normalized (v2)

    def test_was_passthrough_int_encoding(self) -> None:
        rec = _sample_record(was_passthrough=True, filter_name=None)
        assert rec.to_dict()["was_passthrough"] == 1

    def test_recorded_at_auto_populated(self) -> None:
        rec = _sample_record()
        # Should be an ISO 8601 string
        dt = datetime.fromisoformat(rec.recorded_at)
        assert dt.year >= 2026

    def test_posix_path_no_backslash(self) -> None:
        # Even on Windows the path must use forward slashes
        rec = _sample_record(project_path="C:/Users/priya/project")
        assert "\\" not in rec.project_path


# ---------------------------------------------------------------------------
# TrackingDB — SQLite writes
# ---------------------------------------------------------------------------


class TestTrackingDbSqlite:
    def test_schema_created(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        db.flush()
        db.close()
        with sqlite3.connect(str(db_path)) as conn:
            tables = {r[0] for r in conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table'"
            ).fetchall()}
        assert "invocations" in tables
        assert "schema_migrations" in tables

    def test_wal_mode(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        db.flush()
        db.close()
        with sqlite3.connect(str(db_path)) as conn:
            mode = conn.execute("PRAGMA journal_mode").fetchone()[0]
        assert mode == "wal"

    def test_transient_lock_during_schema_init_is_retried(self, tmp_path: Path) -> None:
        """Regression test: two TrackingDB instances initializing against the
        same fresh database at nearly the same moment can transiently collide
        on schema/cleanup, not just the WAL PRAGMA (which was the only
        statement retried before this fix). A losing writer must retry and
        still end up with a working connection, not silently drop every
        record for the rest of its life."""
        db_path = tmp_path / "quor.db"
        call_count = 0
        real_apply_schema = TrackingDB._apply_schema

        def flaky_apply_schema(self: TrackingDB, conn: sqlite3.Connection) -> None:
            nonlocal call_count
            call_count += 1
            if call_count == 1:
                raise sqlite3.OperationalError("database is locked")
            real_apply_schema(self, conn)

        with patch.object(TrackingDB, "_apply_schema", flaky_apply_schema):
            db = TrackingDB(db_path=db_path)
            db.record(_sample_record())
            db.flush(timeout=5.0)
            db.close()

        assert call_count >= 2, "expected at least one retry after the induced lock"
        conn = sqlite3.connect(str(db_path))
        try:
            count = conn.execute("SELECT COUNT(*) FROM invocations").fetchone()[0]
        finally:
            conn.close()
        assert count == 1

    def test_exhausted_schema_init_retries_closes_connection_not_leak_lock(
        self, tmp_path: Path
    ) -> None:
        """Regression test: if schema/cleanup initialization fails on every
        retry (a persistent, not transient, lock), the worker thread must
        warn and exit cleanly instead of crashing uncaught — and, critically,
        must close its half-initialized connection rather than leaving it for
        GC to eventually release. Asserted via the real-world symptom (a
        fresh connection opened immediately after must not be locked out),
        not via an internal call-count — this is the exact bug that caused
        `TestConcurrentWrites.test_two_concurrent_writers_no_data_loss` to
        intermittently fail in CI with `database is locked` on an unrelated,
        separate read connection."""
        db_path = tmp_path / "quor.db"

        def always_fails(self: TrackingDB, conn: sqlite3.Connection) -> None:
            raise sqlite3.OperationalError("database is locked")

        with (
            patch.object(TrackingDB, "_apply_schema", always_fails),
            warnings.catch_warnings(record=True) as caught,
        ):
            warnings.simplefilter("always")
            db = TrackingDB(db_path=db_path)
            db.record(_sample_record())
            db.flush(timeout=5.0)
            db.close()

        assert any("tracking DB unavailable" in str(w.message) for w in caught)

        # The real-world symptom this fix closes: a brand-new connection,
        # opened right after, must not find the database locked.
        conn = sqlite3.connect(str(db_path), timeout=1.0)
        try:
            conn.execute("PRAGMA journal_mode=WAL")
        finally:
            conn.close()

    def test_migration_row_inserted(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        db.flush()
        db.close()
        with sqlite3.connect(str(db_path)) as conn:
            versions = [r[0] for r in conn.execute(
                "SELECT version FROM schema_migrations"
            ).fetchall()]
        assert 2 in versions  # bumped for project_key_normalized (v2)

    def test_record_written_to_sqlite(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        rec = _sample_record()
        db.record(rec)
        db.flush()
        db.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()

        assert row is not None
        assert row["command"] == "git status"
        assert row["original_tokens"] == 100
        assert row["final_tokens"] == 20
        assert row["filter_name"] == "git-status"
        assert row["was_passthrough"] == 0

    def test_path_no_backslash_on_windows(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        db.record(_sample_record(project_path="C:/Users/user/project"))
        db.flush()
        db.close()

        with sqlite3.connect(str(db_path)) as conn:
            stored = conn.execute(
                "SELECT project_path FROM invocations LIMIT 1"
            ).fetchone()[0]

        assert "\\" not in stored
        assert stored == "C:/Users/user/project"

    def test_multiple_records(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        for i in range(5):
            db.record(_sample_record(command=f"cmd {i}"))
        db.flush()
        db.close()

        with sqlite3.connect(str(db_path)) as conn:
            count = conn.execute("SELECT COUNT(*) FROM invocations").fetchone()[0]
        assert count == 5

    def test_passthrough_record(self, tmp_path: Path) -> None:
        db, db_path, _ = _db_and_jsonl(tmp_path)
        db.record(_sample_record(was_passthrough=True, filter_name=None))
        db.flush()
        db.close()

        with sqlite3.connect(str(db_path)) as conn:
            row = conn.execute("SELECT was_passthrough, filter_name FROM invocations").fetchone()
        assert row[0] == 1
        assert row[1] is None

    def test_nonblocking_record(self, tmp_path: Path) -> None:
        """record() must return quickly even under load."""
        db, _, _ = _db_and_jsonl(tmp_path)
        start = time.monotonic()
        for _ in range(50):
            db.record(_sample_record())
        elapsed = time.monotonic() - start
        db.close()
        # Enqueuing 50 records should take <50ms
        assert elapsed < 0.05


# ---------------------------------------------------------------------------
# TrackingDB — JSONL writes
# ---------------------------------------------------------------------------


class TestTrackingDbJsonl:
    def test_record_written_to_jsonl(self, tmp_path: Path) -> None:
        db, _, jsonl_path = _db_and_jsonl(tmp_path)
        rec = _sample_record()
        db.record(rec)
        db.flush()
        db.close()

        lines = jsonl_path.read_bytes().splitlines()
        assert len(lines) == 1
        parsed = orjson.loads(lines[0])
        assert parsed["command"] == "git status"
        assert parsed["filter_name"] == "git-status"

    def test_multiple_records_multiple_lines(self, tmp_path: Path) -> None:
        db, _, jsonl_path = _db_and_jsonl(tmp_path)
        for i in range(3):
            db.record(_sample_record(command=f"cmd {i}"))
        db.flush()
        db.close()

        lines = jsonl_path.read_bytes().splitlines()
        assert len(lines) == 3

    def test_jsonl_fields_match_sqlite_schema(self, tmp_path: Path) -> None:
        db, _db_path, jsonl_path = _db_and_jsonl(tmp_path)
        db.record(_sample_record())
        db.flush()
        db.close()

        parsed = orjson.loads(jsonl_path.read_bytes().splitlines()[0])
        expected_keys = {
            "command", "project_path", "original_tokens", "final_tokens",
            "filter_name", "was_passthrough", "duration_ms", "recorded_at",
            "schema_version",
        }
        assert expected_keys == set(parsed.keys())

    def test_no_jsonl_path_skips_jsonl(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        db = TrackingDB(db_path=db_path, jsonl_path=None)
        db.record(_sample_record())
        db.flush()
        db.close()
        # Should not raise and no jsonl file created
        assert not (tmp_path / "invocations.jsonl").exists()

    def test_write_jsonl_raises_if_called_without_path(self, tmp_path: Path) -> None:
        """TD-002: this was an `assert`, which `python -O` strips silently.
        Calling `_write_jsonl` directly (bypassing `_worker`'s `is not None`
        guard) must raise a real, non-optimizable error."""
        db_path = tmp_path / "quor.db"
        db = TrackingDB(db_path=db_path, jsonl_path=None)
        try:
            with pytest.raises(RuntimeError, match="jsonl_path"):
                db._write_jsonl(_sample_record())
        finally:
            db.close()


# ---------------------------------------------------------------------------
# TrackingDB — 90-day cleanup
# ---------------------------------------------------------------------------


class TestCleanup:
    def _insert_old_record(self, db_path: Path, days_ago: int) -> None:
        old_date = (
            datetime.now(UTC) - timedelta(days=days_ago)
        ).isoformat(timespec="seconds")
        with sqlite3.connect(str(db_path)) as conn:
            conn.execute(
                """INSERT INTO invocations
                   (command, project_path, original_tokens, final_tokens,
                    filter_name, was_passthrough, duration_ms, recorded_at, schema_version)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                ("old cmd", "/project", 10, 5, "git", 0, 1.0, old_date, 1),
            )
            conn.commit()

    def test_old_records_removed(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        # First create the schema by opening a DB
        db = TrackingDB(db_path=db_path)
        db.flush()
        db.close()

        # Insert a record that is 100 days old (past the 90-day window)
        self._insert_old_record(db_path, days_ago=100)

        with sqlite3.connect(str(db_path)) as conn:
            count_before = conn.execute("SELECT COUNT(*) FROM invocations").fetchone()[0]
        assert count_before == 1

        # Re-open DB — cleanup runs on connect
        db2 = TrackingDB(db_path=db_path)
        db2.flush()
        db2.close()

        with sqlite3.connect(str(db_path)) as conn:
            count_after = conn.execute("SELECT COUNT(*) FROM invocations").fetchone()[0]
        assert count_after == 0

    def test_recent_records_preserved(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        db = TrackingDB(db_path=db_path)
        db.flush()
        db.close()

        # Insert a 10-day-old record (within the 90-day window)
        self._insert_old_record(db_path, days_ago=10)

        db2 = TrackingDB(db_path=db_path)
        db2.flush()
        db2.close()

        with sqlite3.connect(str(db_path)) as conn:
            count = conn.execute("SELECT COUNT(*) FROM invocations").fetchone()[0]
        assert count == 1


# ---------------------------------------------------------------------------
# query_gain
# ---------------------------------------------------------------------------


class TestQueryGain:
    def _populate(self, db_path: Path, records: list[dict]) -> None:
        with sqlite3.connect(str(db_path)) as conn:
            # Create schema
            schema_sql = (
                Path(__file__).parent.parent.parent
                / "quor" / "tracking" / "schema.sql"
            ).read_text(encoding="utf-8")
            conn.executescript(schema_sql)
            for r in records:
                conn.execute(
                    """INSERT INTO invocations
                       (command, project_path, original_tokens, final_tokens,
                        filter_name, was_passthrough, duration_ms, recorded_at, schema_version)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        r.get("command", "git status"),
                        r.get("project_path", "/proj"),
                        r.get("original_tokens", 100),
                        r.get("final_tokens", 20),
                        r.get("filter_name", "git"),
                        r.get("was_passthrough", 0),
                        r.get("duration_ms", 10.0),
                        r.get("recorded_at", datetime.now(UTC).isoformat()),
                        1,
                    ),
                )
            conn.commit()

    def test_empty_db_returns_zeros(self, tmp_path: Path) -> None:
        report = query_gain(tmp_path / "missing.db", tmp_path)
        assert report.total_invocations == 0
        assert report.tokens_saved == 0
        assert report.tokens_before == 0
        assert report.tokens_after == 0
        assert report.filter_hit_rate == 0.0

    def test_basic_totals(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},
            {"original_tokens": 200, "final_tokens": 50, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.total_invocations == 2
        assert report.tokens_saved == (100 - 20) + (200 - 50)  # 230

    def test_tokens_before_and_after(self, tmp_path: Path) -> None:
        """tokens_before/tokens_after are the raw sums driving tokens_saved —
        this must hold exactly (tokens_before - tokens_after == tokens_saved),
        since they're read from the same original_tokens/final_tokens columns."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},
            {"original_tokens": 200, "final_tokens": 50, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.tokens_before == 300
        assert report.tokens_after == 70
        assert report.tokens_before - report.tokens_after == report.tokens_saved

    def test_gross_savings_and_overhead_decomposition(self, tmp_path: Path) -> None:
        """QB-017 gain hardening: gross_savings/gross_overhead split the same
        per-row difference tokens_saved already sums, into positive and
        negative parts. Three rows: one genuinely compressed, one genuinely
        grew (the QB-017 tee-overhead shape), one unchanged."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},  # -80
            {"original_tokens": 21, "final_tokens": 43, "project_path": "/proj"},   # +22
            {"original_tokens": 50, "final_tokens": 50, "project_path": "/proj"},   # 0
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.gross_savings == 80
        assert report.gross_overhead == 22
        assert report.negative_row_count == 1
        # Exact identity: the decomposition must always net back out to the
        # same tokens_saved figure already computed the original way.
        assert report.gross_savings - report.gross_overhead == report.tokens_saved
        assert report.tokens_saved == (100 - 20) + (21 - 43) + (50 - 50)  # 58

    def test_gross_savings_and_overhead_all_positive(self, tmp_path: Path) -> None:
        """When nothing grew, gross_overhead and negative_row_count are both
        zero — gross_savings equals tokens_saved exactly."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},
            {"original_tokens": 200, "final_tokens": 50, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.gross_overhead == 0
        assert report.negative_row_count == 0
        assert report.gross_savings == report.tokens_saved == 230

    def test_gross_savings_and_overhead_all_negative(self, tmp_path: Path) -> None:
        """When every row grew, gross_savings is zero and gross_overhead
        equals the (negative) tokens_saved's magnitude exactly."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"original_tokens": 21, "final_tokens": 43, "project_path": "/proj"},
            {"original_tokens": 10, "final_tokens": 15, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.gross_savings == 0
        assert report.negative_row_count == 2
        assert report.gross_overhead == 27
        assert report.tokens_saved == -27

    def test_empty_db_gross_fields_are_zero(self, tmp_path: Path) -> None:
        report = query_gain(tmp_path / "missing.db", tmp_path)
        assert report.gross_savings == 0
        assert report.gross_overhead == 0
        assert report.negative_row_count == 0

    def test_glob_project_scoping(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/my/project", "original_tokens": 100, "final_tokens": 10},
            {"project_path": "/other/project", "original_tokens": 200, "final_tokens": 50},
        ])
        report = query_gain(db_path, Path("/my/project"))
        assert report.total_invocations == 1
        assert report.tokens_saved == 90

    def test_passthrough_counted(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"was_passthrough": 1, "filter_name": None, "project_path": "/proj"},
            {"was_passthrough": 0, "filter_name": "git", "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.passthrough_count == 1
        assert abs(report.filter_hit_rate - 0.5) < 1e-9

    def test_top_filters_ordered_by_savings(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"filter_name": "pytest", "original_tokens": 500, "final_tokens": 50,
             "project_path": "/proj"},
            {"filter_name": "git", "original_tokens": 100, "final_tokens": 90,
             "project_path": "/proj"},
            {"filter_name": "pytest", "original_tokens": 300, "final_tokens": 30,
             "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.top_filters[0][0] == "pytest"
        assert report.top_filters[0][1] == (500 - 50) + (300 - 30)  # 720

    def test_days_filter(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        old_date = (
            datetime.now(UTC) - timedelta(days=60)
        ).isoformat(timespec="seconds")
        self._populate(db_path, [
            {"recorded_at": old_date, "project_path": "/proj"},
        ])
        # 30-day window should exclude the 60-day-old record
        report = query_gain(db_path, Path("/proj"), days=30)
        assert report.total_invocations == 0

    def test_glob_includes_subdirectories(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/proj"},
            {"project_path": "/proj/subdir"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.total_invocations == 2

    def test_project_identity_case_insensitive_consistency(self, tmp_path: Path) -> None:
        """'C:/Workspace' and 'c:/workspace' must be treated as one project —
        case differences anywhere in the path (not just the drive letter)
        must not split one project's history into two logical entries."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "C:/Workspace", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "c:/workspace", "original_tokens": 200, "final_tokens": 50},
        ])
        report = query_gain(db_path, Path("C:/Workspace"))
        assert report.total_invocations == 2
        assert report.tokens_before == 300
        assert report.tokens_after == 70

    def test_project_identity_no_sibling_leakage(self, tmp_path: Path) -> None:
        """A sibling directory that merely shares a text prefix — with no
        path-separator boundary — must never be included. Regression for
        the naive `GLOB "{project}*"` matching "/workspace-other" under a
        query for "/workspace"."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/workspace", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "/workspace-other", "original_tokens": 999, "final_tokens": 999},
        ])
        report = query_gain(db_path, Path("/workspace"))
        assert report.total_invocations == 1
        assert report.tokens_before == 100
        assert report.tokens_after == 20

    def test_project_identity_subdirectory_inclusion(self, tmp_path: Path) -> None:
        """True subdirectories, separated by "/", must still be included —
        the same canonical rule that excludes sibling-prefix leakage above
        must not also exclude genuine nesting."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/workspace", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "/workspace/subdir", "original_tokens": 50, "final_tokens": 10},
        ])
        report = query_gain(db_path, Path("/workspace"))
        assert report.total_invocations == 2
        assert report.tokens_before == 150
        assert report.tokens_after == 30

    def test_project_identity_glob_special_characters_in_path(self, tmp_path: Path) -> None:
        """A real directory name containing literal brackets — special to
        the old GLOB-based design, requiring escaping there — must still
        match correctly now that matching is LIKE-based: brackets have no
        special meaning to LIKE at all, so this passes without any
        escaping being needed for this particular character set. Kept as
        the same named regression test the GLOB design required, now
        verifying the LIKE-based replacement handles it too."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "c:/my[client]project", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "c:/my[client]project/subdir", "original_tokens": 50, "final_tokens": 10},
        ])
        report = query_gain(db_path, Path("c:/my[client]project"))
        assert report.total_invocations == 2
        assert report.tokens_before == 150
        assert report.tokens_after == 30

    def test_project_identity_like_wildcard_characters_do_not_overmatch(self, tmp_path: Path) -> None:
        """'_' is LIKE's "match any single character" wildcard — and
        extremely common in real directory names ("my_project"). Without
        escaping, querying "my_project" would have its subdirectory pattern
        "my_project/%" also match "myXproject/subdir" for any character X.
        This must not happen."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "c:/my_project", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "c:/myXproject/subdir", "original_tokens": 999, "final_tokens": 999},
        ])
        report = query_gain(db_path, Path("c:/my_project"))
        assert report.total_invocations == 1
        assert report.tokens_before == 100
        assert report.tokens_after == 20

    def test_degenerate_root_key_rejected(self, tmp_path: Path) -> None:
        """A normalized key with no directory segment of its own (empty, or
        a bare drive letter) must be rejected outright rather than silently
        turned into a match-everything wildcard."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "C:/alpha-project"},
            {"project_path": "C:/beta-project"},
        ])
        for degenerate in ("/", "C:/", "c:"):
            with pytest.raises(ValueError, match="too broad"):
                query_gain(db_path, Path(degenerate))

    def test_drive_root_query_no_longer_overmatches(self, tmp_path: Path) -> None:
        """Regression for the demonstrated wildcard-explosion bug: querying
        a bare drive root must raise rather than silently sweeping in every
        unrelated project on that drive."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "C:/alpha-project"},
            {"project_path": "C:/beta-project"},
            {"project_path": "C:/gamma-project"},
        ])
        with pytest.raises(ValueError):
            query_gain(db_path, Path("C:/"))

    def test_project_identity_deeply_nested_subdirectories(self, tmp_path: Path) -> None:
        """QB-017 hardening: the subdirectory LIKE pattern ("{key}/%") must
        match arbitrarily deep nesting, not just one level down — '%' spans
        multiple path segments, so a 3-level-deep subdirectory must be
        included exactly like a 1-level one already is
        (test_project_identity_subdirectory_inclusion)."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/workspace", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "/workspace/a/b/c", "original_tokens": 50, "final_tokens": 10},
            {"project_path": "/workspace/a/b/c/d/e", "original_tokens": 40, "final_tokens": 5},
        ])
        report = query_gain(db_path, Path("/workspace"))
        assert report.total_invocations == 3
        assert report.tokens_before == 190
        assert report.tokens_after == 35

    def test_project_identity_case_insensitive_sibling_exclusion(self, tmp_path: Path) -> None:
        """Case-insensitivity and sibling-prefix exclusion must compose: a
        sibling whose only relation to the queried project is a shared text
        prefix must still be excluded even when its casing differs from
        both the query and the stored separator boundary rule."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/Workspace", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "/WORKSPACE-OTHER", "original_tokens": 999, "final_tokens": 999},
        ])
        report = query_gain(db_path, Path("/workspace"))
        assert report.total_invocations == 1
        assert report.tokens_before == 100
        assert report.tokens_after == 20

    def test_project_identity_case_insensitive_subdirectory_inclusion(
        self, tmp_path: Path
    ) -> None:
        """Case-insensitivity and subdirectory inclusion must also compose:
        a genuine subdirectory recorded with different casing than the
        query must still be included."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/Workspace", "original_tokens": 100, "final_tokens": 20},
            {"project_path": "/WORKSPACE/Subdir", "original_tokens": 50, "final_tokens": 10},
        ])
        report = query_gain(db_path, Path("/workspace"))
        assert report.total_invocations == 2
        assert report.tokens_before == 150
        assert report.tokens_after == 30

    def test_project_identity_trailing_slash_on_query_path(self, tmp_path: Path) -> None:
        """Integration-level companion to TestNormalizeProjectPath's
        test_trailing_slash_insensitive: a query Path constructed with a
        trailing slash must resolve to the same project_key as one without,
        end to end through query_gain — not just at the normalize function
        in isolation."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"project_path": "/workspace", "original_tokens": 100, "final_tokens": 20},
        ])
        report = query_gain(db_path, Path("/workspace/"))
        assert report.total_invocations == 1
        assert report.tokens_before == 100
        assert report.tokens_after == 20

    def test_lazy_backfill_populates_missing_columns_for_pre_v2_database(
        self, tmp_path: Path
    ) -> None:
        """A database created before schema v2 has no project_key_normalized
        column at all (not merely a NULL value in it) — query_gain() must
        add the column and backfill existing rows on first read, with no
        manual migration step, and the row must be included in the result
        exactly as if it had always had it."""
        db_path = tmp_path / "quor.db"
        pre_v2_schema = """
            CREATE TABLE invocations (
                id               INTEGER PRIMARY KEY AUTOINCREMENT,
                command          TEXT    NOT NULL,
                project_path     TEXT    NOT NULL,
                original_tokens  INTEGER NOT NULL DEFAULT 0,
                final_tokens     INTEGER NOT NULL DEFAULT 0,
                filter_name      TEXT,
                was_passthrough  INTEGER NOT NULL DEFAULT 0,
                duration_ms      REAL    NOT NULL DEFAULT 0,
                recorded_at      TEXT    NOT NULL DEFAULT (datetime('now')),
                schema_version   INTEGER NOT NULL DEFAULT 1
            );
        """
        with sqlite3.connect(str(db_path)) as conn:
            conn.executescript(pre_v2_schema)
            conn.execute(
                """INSERT INTO invocations
                   (command, project_path, original_tokens, final_tokens,
                    filter_name, was_passthrough, duration_ms, recorded_at, schema_version)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    "git status", "C:/legacy-project", 100, 20,
                    "git-status", 0, 1.0, datetime.now(UTC).isoformat(), 1,
                ),
            )
            conn.commit()

        report = query_gain(db_path, Path("C:/legacy-project"))
        assert report.total_invocations == 1
        assert report.tokens_before == 100
        assert report.tokens_after == 20

        # Confirm the row was actually backfilled, not found some other way.
        with sqlite3.connect(str(db_path)) as conn:
            row = conn.execute(
                "SELECT project_key_normalized FROM invocations"
            ).fetchone()
        assert row == ("c:/legacy-project",)

    def test_read_hook_invocations_zero_when_only_bash_rows(self, tmp_path: Path) -> None:
        """A project/window with only Bash-produced rows (no command
        starting with the literal "Read: " prefix claude_read.py always
        uses) must report read_hook_invocations == 0 — the signal `quor
        gain`'s new notice checks before telling a user Read-hook features
        aren't represented."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"command": "git status", "original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},
            {"command": "cat file.py", "original_tokens": 200, "final_tokens": 50, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.total_invocations == 2
        assert report.read_hook_invocations == 0

    def test_read_hook_invocations_counts_only_read_prefixed_rows(self, tmp_path: Path) -> None:
        """A mix of Bash and Read rows: read_hook_invocations counts exactly
        the "Read: " ones, independent of total_invocations."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"command": "git status", "original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},
            {"command": "Read: notes.md", "original_tokens": 500, "final_tokens": 300, "project_path": "/proj"},
            {"command": "Read: app.py", "original_tokens": 400, "final_tokens": 250, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.total_invocations == 3
        assert report.read_hook_invocations == 2

    def test_read_hook_invocations_scoped_by_project_like_everything_else(self, tmp_path: Path) -> None:
        """A Read row in a *different* project must not count toward this
        project's read_hook_invocations — same project scoping every other
        GainReport field already respects."""
        db_path = tmp_path / "quor.db"
        self._populate(db_path, [
            {"command": "Read: notes.md", "original_tokens": 500, "final_tokens": 300, "project_path": "/other-proj"},
            {"command": "git status", "original_tokens": 100, "final_tokens": 20, "project_path": "/proj"},
        ])
        report = query_gain(db_path, Path("/proj"))
        assert report.total_invocations == 1
        assert report.read_hook_invocations == 0

    def test_read_hook_invocations_zero_on_empty_db(self, tmp_path: Path) -> None:
        report = query_gain(tmp_path / "missing.db", tmp_path)
        assert report.read_hook_invocations == 0


class TestNormalizeProjectPath:
    """Direct unit coverage for the single canonical project-identity rule
    (query_gain's matching is only as correct as this function's contract)."""

    def test_case_insensitive(self) -> None:
        assert normalize_project_path("C:/Workspace") == normalize_project_path("c:/workspace")

    @pytest.mark.skipif(
        os.name != "nt",
        reason="Path(...) only parses backslash as a separator on Windows "
        "(WindowsPath); on POSIX, Path is PosixPath and treats backslash as "
        "a literal character within one path component, so this input "
        "isn't a backslash-separated path at all off Windows.",
    )
    def test_backslashes_normalized_to_posix(self) -> None:
        assert normalize_project_path(Path("C:\\Users\\dev\\project")) == normalize_project_path(
            "C:/Users/dev/project"
        )

    def test_trailing_slash_insensitive(self) -> None:
        assert normalize_project_path("/proj/") == normalize_project_path("/proj")

    def test_accepts_str_or_path_identically(self) -> None:
        assert normalize_project_path(Path("/proj")) == normalize_project_path("/proj")

    def test_distinct_projects_remain_distinct(self) -> None:
        assert normalize_project_path("/workspace") != normalize_project_path("/workspace-other")

    def test_normalize_project_path_idempotence(self) -> None:
        """Re-normalizing an already-normalized value must be a no-op — the
        SQL layer relies on project_key being stable and final by the time
        it reaches a query parameter, with no further transformation needed
        or applied anywhere else."""
        once = normalize_project_path("C:\\Users\\dev\\project/")
        twice = normalize_project_path(once)
        assert once == twice


# ---------------------------------------------------------------------------
# get_tracking_db — factory (uses platformdirs, already patched by conftest)
# ---------------------------------------------------------------------------


class TestGetTrackingDb:
    def test_returns_tracking_db(self) -> None:
        db = get_tracking_db()
        assert isinstance(db, TrackingDB)
        db.close()

    def test_db_in_data_dir(self) -> None:
        import platformdirs

        db = get_tracking_db()
        data_dir = Path(platformdirs.user_data_dir("quor"))
        db.close()
        # DB file should be created inside the platformdirs data dir
        assert (data_dir / "quor.db").exists()

    def test_tracking_failure_does_not_raise(self, tmp_path: Path) -> None:
        """A write failure in the background worker must not propagate to the caller."""
        db = TrackingDB(db_path=tmp_path / "quor.db")
        with (
            patch.object(db, "_write_sqlite", side_effect=RuntimeError("disk full")),
            pytest.warns(UserWarning, match="tracking write error"),
        ):
            db.record(_sample_record(command="should not crash"))
            db.flush()
        db.close()


# ---------------------------------------------------------------------------
# Dispatcher integration with tracking
# ---------------------------------------------------------------------------


class TestDispatcherTracking:
    def test_record_written_on_filtered_dispatch(self, tmp_path: Path) -> None:
        db_path = tmp_path / "quor.db"
        jsonl_path = tmp_path / "invocations.jsonl"
        tracking = TrackingDB(db_path=db_path, jsonl_path=jsonl_path)

        proc = MagicMock(spec=subprocess.CompletedProcess)
        proc.stdout = (
            "FAILED tests/test_x.py::test_y\n"
            "    AssertionError: got False\n"
        )
        proc.returncode = 1

        captured_stdout = io.StringIO()
        with (
            patch("subprocess.run", return_value=proc),
            patch("sys.stdout", captured_stdout),
        ):
            from quor.adapters.dispatcher import run_dispatch
            run_dispatch(["pytest", "tests/"], tracking=tracking)

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            row = conn.execute(
                "SELECT command, filter_name, was_passthrough FROM invocations LIMIT 1"
            ).fetchone()
        assert row is not None
        assert row[0] == "pytest tests/"
        assert row[2] == 0  # not passthrough

    def test_passthrough_recorded(self, tmp_path: Path) -> None:
        """When registry.find() returns None the record is marked passthrough."""
        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)

        proc = MagicMock(spec=subprocess.CompletedProcess)
        proc.stdout = "hello world\n"
        proc.returncode = 0

        captured_stdout = io.StringIO()
        with (
            patch("subprocess.run", return_value=proc),
            patch("sys.stdout", captured_stdout),
            patch("quor.adapters.dispatcher.FilterRegistry") as mock_reg_cls,
        ):
            mock_reg_cls.return_value.find.return_value = None
            from quor.adapters.dispatcher import run_dispatch
            run_dispatch(["echo", "hello"], tracking=tracking)

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            row = conn.execute(
                "SELECT was_passthrough, filter_name FROM invocations LIMIT 1"
            ).fetchone()
        assert row is not None
        assert row[0] == 1       # passthrough
        assert row[1] is None    # no filter


# ---------------------------------------------------------------------------
# Read hook integration with tracking (QB-007D) — Read becomes another
# producer of InvocationRecord via the same track_invocation() helper
# TestDispatcherTracking above exercises through run_dispatch(). These tests
# call quor.adapters.claude_read._compress_read_output() directly, the same
# way TestDispatcherTracking calls run_dispatch() directly, and then read
# the result back out of the real SQLite/JSONL stores — no Read-specific
# storage or aggregation exists anywhere in this path.
# ---------------------------------------------------------------------------


def _read_hook_input(file_path: str, tool_response: object) -> Any:
    from quor.adapters.base import PostToolUseHookInput

    return PostToolUseHookInput.model_validate(
        {
            "tool_name": "Read",
            "tool_input": {"file_path": file_path},
            "tool_response": tool_response,
        }
    )


class TestReadTracking:
    _hook_input = staticmethod(_read_hook_input)

    def test_compressed_read_tracked(self, tmp_path: Path) -> None:
        """An oversized markdown document that genuinely compresses is
        tracked with the matched filter name and was_passthrough=False,
        exactly like a compressed Bash invocation."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        large_content = "line of filler text. " * 10_000
        hook_input = self._hook_input("notes.md", large_content)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None
        assert len(result) < len(large_content)

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["command"] == "Read: notes.md"
        assert row["filter_name"] == "markdown"
        assert row["was_passthrough"] == 0
        assert row["original_tokens"] > row["final_tokens"]

    def test_unchanged_read_tracked(self, tmp_path: Path) -> None:
        """A small markdown document under the filter's compression budget
        is tracked (filter matched, was_passthrough=False) even though
        updatedToolOutput ends up omitted because nothing changed."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        small_content = "# Heading\n\nBody text.\n"
        hook_input = self._hook_input("notes.md", small_content)

        result = _compress_read_output(hook_input, tracking)
        assert result is None  # nothing to report — content unchanged

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["command"] == "Read: notes.md"
        assert row["filter_name"] == "markdown"
        assert row["was_passthrough"] == 0
        assert row["original_tokens"] == row["final_tokens"]

    def test_unsupported_file_type_tracked(self, tmp_path: Path) -> None:
        """A file type outside both the Read allowlist and the QB-005F
        source-code extension mapping (here: a .json file, which
        FilterRegistry still routes to the Bash-oriented `generic` filter)
        is tracked as a passthrough — no document/source-code filter is
        genuinely applied to it, matching how dispatcher tracks a Bash
        command that matched no filter at all. (.py was this test's example
        prior to QB-005F; it is now a genuinely supported source-code
        extension — see TestReadSourceCodeTracking below.)"""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        hook_input = self._hook_input("data.json", '{"hello": "world"}\n')

        result = _compress_read_output(hook_input, tracking)
        assert result is None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["command"] == "Read: data.json"
        assert row["filter_name"] is None
        assert row["was_passthrough"] == 1

    def test_no_matching_filter_tracked(self, tmp_path: Path) -> None:
        """A file path FilterRegistry cannot route at all (whitespace in the
        path defeats every built-in filter's whitespace-free anchor) is
        also tracked as a passthrough."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        hook_input = self._hook_input("My Documents/notes.md", "content\n")

        result = _compress_read_output(hook_input, tracking)
        assert result is None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute(
                "SELECT filter_name, was_passthrough FROM invocations LIMIT 1"
            ).fetchone()
        assert row is not None
        assert row["was_passthrough"] == 1

    def test_tracking_failure_stays_fail_open(self, tmp_path: Path) -> None:
        """A tracking write failure must never affect what the Read hook
        returns — mirrors dispatcher's own tracking fail-open guarantee."""
        from quor.adapters.claude_read import _compress_read_output

        tracking = TrackingDB(db_path=tmp_path / "quor.db")
        large_content = "line of filler text. " * 10_000
        hook_input = self._hook_input("notes.md", large_content)

        with (
            patch.object(tracking, "record", side_effect=RuntimeError("disk full")),
            pytest.warns(UserWarning, match="tracking record error"),
        ):
            result = _compress_read_output(hook_input, tracking)

        tracking.close()
        assert result is not None
        assert len(result) < len(large_content)

    def test_jsonl_fallback_for_read(self, tmp_path: Path) -> None:
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        jsonl_path = tmp_path / "invocations.jsonl"
        tracking = TrackingDB(db_path=db_path, jsonl_path=jsonl_path)
        hook_input = self._hook_input("notes.md", "# Heading\n\nBody text.\n")

        _compress_read_output(hook_input, tracking)
        tracking.flush()
        tracking.close()

        lines = jsonl_path.read_bytes().splitlines()
        assert len(lines) == 1
        parsed = orjson.loads(lines[0])
        assert parsed["command"] == "Read: notes.md"
        assert parsed["filter_name"] == "markdown"

    def test_project_identity_matches_bash_rows(self, tmp_path: Path) -> None:
        """A Read row and a Bash row recorded for the same project must
        resolve to the same project_key_normalized — Read reuses
        Path.cwd().as_posix(), the identical project-resolution rule
        dispatcher.py already uses, with no Read-specific logic."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)

        with patch("pathlib.Path.cwd", return_value=tmp_path):
            hook_input = self._hook_input("notes.md", "# Heading\n\nBody text.\n")
            _compress_read_output(hook_input, tracking)
            tracking.record(_sample_record(project_path=tmp_path.as_posix()))

        tracking.flush()
        tracking.close()

        report = query_gain(db_path, tmp_path)
        assert report.total_invocations == 2

    def test_multiple_read_operations_aggregate(self, tmp_path: Path) -> None:
        """Several Read calls in the same project sum into one project's
        totals via the exact same query_gain() aggregation Bash rows use —
        no Read-specific aggregation path exists."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        large_content = "line of filler text. " * 10_000

        with patch("pathlib.Path.cwd", return_value=tmp_path):
            for name in ("a.md", "b.md", "c.txt"):
                hook_input = self._hook_input(name, large_content)
                _compress_read_output(hook_input, tracking)

        tracking.flush()
        tracking.close()

        report = query_gain(db_path, tmp_path)
        assert report.total_invocations == 3
        assert report.tokens_saved > 0
        assert report.passthrough_count == 0

    # -----------------------------------------------------------------
    # DOCX/PDF extraction tracking (QB-007E4) — no schema/tracking-side
    # change was needed: track_invocation() is called exactly the same
    # way for the extraction path as for the direct markdown/text path
    # above; these tests exist to prove that in practice, not because the
    # tracking code itself differs.
    # -----------------------------------------------------------------

    def test_docx_extraction_and_compression_tracked(self, tmp_path: Path) -> None:
        """original_tokens reflects the raw tool_response (the pre-
        extraction Read result); final_tokens reflects the actually-
        returned, extracted-and-compressed Markdown — not any
        intermediate, extraction-only value."""
        import docx

        from quor.adapters.claude_read import _compress_read_output

        d = docx.Document()
        d.add_heading("Design Notes", level=1)
        d.add_paragraph("REQ-1: must survive extraction and compression.")
        for _ in range(150):
            d.add_paragraph("This is an ordinary sentence of filler prose, repeated.")
        docx_path = tmp_path / "report.docx"
        d.save(str(docx_path))

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        raw_tool_response = "<binary Read result placeholder>"
        hook_input = self._hook_input(str(docx_path), raw_tool_response)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["command"] == f"Read: {docx_path}"
        assert row["filter_name"] == "markdown"
        assert row["was_passthrough"] == 0
        assert row["original_tokens"] == count_tokens(raw_tool_response)
        assert row["final_tokens"] == count_tokens(result)

    def test_pdf_extraction_and_compression_tracked(self, tmp_path: Path) -> None:
        from reportlab.pdfgen import canvas

        from quor.adapters.claude_read import _compress_read_output

        pdf_path = tmp_path / "report.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=(500, 3000))
        c.setFont("Helvetica-Bold", 20)
        c.drawString(72, 2950, "Design Notes")
        c.setFont("Helvetica", 11)
        y = 2910
        for _ in range(150):
            c.drawString(72, y, "This is an ordinary sentence of filler prose, repeated.")
            y -= 16
        c.save()

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        raw_tool_response = "<binary Read result placeholder>"
        hook_input = self._hook_input(str(pdf_path), raw_tool_response)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["filter_name"] == "markdown"
        assert row["was_passthrough"] == 0
        assert row["original_tokens"] == count_tokens(raw_tool_response)
        assert row["final_tokens"] == count_tokens(result)

    def test_extraction_failure_tracked_as_passthrough(self, tmp_path: Path) -> None:
        """A .docx path that fails to extract (here: doesn't exist on disk)
        is tracked exactly like "no filter matched" — filter_name=None,
        was_passthrough=True — not as a special extraction-specific
        outcome."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        missing_path = tmp_path / "does_not_exist.docx"
        raw_tool_response = "original content"
        hook_input = self._hook_input(str(missing_path), raw_tool_response)

        with warnings.catch_warnings():
            warnings.simplefilter("ignore")
            result = _compress_read_output(hook_input, tracking)
        assert result is None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["filter_name"] is None
        assert row["was_passthrough"] == 1
        assert row["original_tokens"] == count_tokens(raw_tool_response)
        assert row["final_tokens"] == count_tokens(raw_tool_response)

    def test_multiple_docx_pdf_reads_aggregate_with_markdown_reads(self, tmp_path: Path) -> None:
        """DOCX/PDF rows aggregate into the same project totals as direct
        .md rows — no Read-format-specific aggregation path exists, same
        guarantee test_multiple_read_operations_aggregate already proves
        for markdown/text."""
        import docx

        from quor.adapters.claude_read import _compress_read_output

        d = docx.Document()
        d.add_heading("Notes", level=1)
        for _ in range(150):
            d.add_paragraph("This is an ordinary sentence of filler prose, repeated.")
        docx_path = tmp_path / "notes.docx"
        d.save(str(docx_path))

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        large_content = "line of filler text. " * 10_000

        with patch("pathlib.Path.cwd", return_value=tmp_path):
            _compress_read_output(
                self._hook_input(str(docx_path), "placeholder"), tracking
            )
            _compress_read_output(
                self._hook_input(str(tmp_path / "a.md"), large_content), tracking
            )

        tracking.flush()
        tracking.close()

        report = query_gain(db_path, tmp_path)
        assert report.total_invocations == 2
        assert report.passthrough_count == 0


# ---------------------------------------------------------------------------
# Source-code Read tracking (QB-005F) — no schema/tracking-side change was
# needed here either: _compress_via_named_filter() (the helper the
# source-code path shares with the DOCX/PDF extraction path above) calls
# track_invocation() exactly the same way. These tests exist to prove that
# in practice for the new by-name source-code lookup, the same way
# TestReadTracking's extraction tests prove it for the by-name document
# lookup.
# ---------------------------------------------------------------------------


class TestReadSourceCodeTracking:
    """Uses the module-level `_read_hook_input` helper `TestReadTracking`
    above also uses — deliberately not a subclass of `TestReadTracking`,
    since inheriting a test class in pytest re-collects and re-runs every
    inherited test method under the subclass too, which is not the intent
    here (only the tiny payload-building helper is shared)."""

    _hook_input = staticmethod(_read_hook_input)

    _PY_SOURCE = (
        "def fetch_data(url):\n"
        '    """Fetch data from a URL."""\n'
        "    response = make_request(url)\n"
        "    return response.json()\n"
    )

    def test_python_read_tracked_with_cat_python_filter(self, tmp_path: Path) -> None:
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        hook_input = self._hook_input("app.py", self._PY_SOURCE)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None
        assert "response = make_request(url)" not in result

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["command"] == "Read: app.py"
        assert row["filter_name"] == "cat-python"
        assert row["was_passthrough"] == 0
        assert row["original_tokens"] == count_tokens(self._PY_SOURCE)
        assert row["final_tokens"] == count_tokens(result)

    def test_javascript_read_tracked_with_cat_javascript_filter(self, tmp_path: Path) -> None:
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        source = (
            "function fetchData(url) {\n"
            "  const response = makeRequest(url);\n"
            "  return response.json();\n"
            "}\n"
        )
        hook_input = self._hook_input("app.js", source)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["filter_name"] == "cat-javascript"
        assert row["was_passthrough"] == 0

    def test_typescript_read_tracked_with_cat_typescript_filter(self, tmp_path: Path) -> None:
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        source = (
            "function fetchData(url: string): Promise<unknown> {\n"
            "  const response = makeRequest(url);\n"
            "  return response.json();\n"
            "}\n"
        )
        hook_input = self._hook_input("app.ts", source)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["filter_name"] == "cat-typescript"
        assert row["was_passthrough"] == 0

    def test_tsx_read_tracked_with_cat_tsx_filter(self, tmp_path: Path) -> None:
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        source = (
            "export function Button({ label }: { label: string }) {\n"
            '  const handleClick = () => { console.log("clicked"); };\n'
            "  return <button onClick={handleClick}>{label}</button>;\n"
            "}\n"
        )
        hook_input = self._hook_input("Button.tsx", source)

        result = _compress_read_output(hook_input, tracking)
        assert result is not None

        tracking.flush()
        tracking.close()

        with sqlite3.connect(str(db_path)) as conn:
            conn.row_factory = sqlite3.Row
            row = conn.execute("SELECT * FROM invocations LIMIT 1").fetchone()
        assert row is not None
        assert row["filter_name"] == "cat-tsx"
        assert row["was_passthrough"] == 0

    def test_multiple_source_code_reads_aggregate_with_markdown_reads(
        self, tmp_path: Path
    ) -> None:
        """Python/JS/TS/TSX rows aggregate into the same project totals as
        direct .md rows — no source-code-specific aggregation path exists,
        same guarantee TestReadTracking already proves for markdown/text and
        DOCX/PDF."""
        from quor.adapters.claude_read import _compress_read_output

        db_path = tmp_path / "quor.db"
        tracking = TrackingDB(db_path=db_path)
        large_content = "line of filler text. " * 10_000

        with patch("pathlib.Path.cwd", return_value=tmp_path):
            _compress_read_output(self._hook_input("app.py", self._PY_SOURCE), tracking)
            _compress_read_output(self._hook_input(str(tmp_path / "a.md"), large_content), tracking)

        tracking.flush()
        tracking.close()

        report = query_gain(db_path, tmp_path)
        assert report.total_invocations == 2
        assert report.passthrough_count == 0


# ---------------------------------------------------------------------------
# Concurrency — WAL mode under two simultaneous writers (P1 item 7)
# ---------------------------------------------------------------------------


class TestConcurrentWrites:
    """Prove WAL mode allows two sessions to write simultaneously without data loss.

    This simulates two Claude Code sessions running commands at the same time and
    both writing their invocation records to the same shared SQLite database.
    Without WAL mode, one writer would get SQLITE_BUSY and records would be lost.
    """

    N_RECORDS = 30  # per writer; total expected = 2 x N_RECORDS

    def _write_records(self, db_path: Path, start: int, count: int) -> None:
        db = TrackingDB(db_path=db_path)
        for i in range(count):
            db.record(
                _sample_record(
                    command=f"git status {start + i}",
                    project_path="/concurrent-project",
                )
            )
        db.flush(timeout=5.0)
        db.close()

    def test_two_concurrent_writers_no_data_loss(self, tmp_path: Path) -> None:
        writer_a = threading.Thread(
            target=self._write_records,
            args=(tmp_path / "quor.db", 0, self.N_RECORDS),
        )
        writer_b = threading.Thread(
            target=self._write_records,
            args=(tmp_path / "quor.db", self.N_RECORDS, self.N_RECORDS),
        )

        writer_a.start()
        writer_b.start()
        writer_a.join(timeout=15)
        writer_b.join(timeout=15)

        assert not writer_a.is_alive(), "Writer A did not complete within timeout"
        assert not writer_b.is_alive(), "Writer B did not complete within timeout"

        with sqlite3.connect(str(tmp_path / "quor.db")) as conn:
            count = conn.execute("SELECT COUNT(*) FROM invocations").fetchone()[0]

        expected = self.N_RECORDS * 2
        assert count == expected, (
            f"Expected {expected} records from two concurrent writers, got {count}. "
            "Data loss detected — WAL mode may not be effective."
        )

    def test_concurrent_writers_wal_mode_confirmed(self, tmp_path: Path) -> None:
        """WAL journal mode must be set even when two writers open the same DB."""
        db_path = tmp_path / "quor.db"

        writer_a = threading.Thread(
            target=self._write_records, args=(db_path, 0, 5)
        )
        writer_b = threading.Thread(
            target=self._write_records, args=(db_path, 5, 5)
        )
        writer_a.start()
        writer_b.start()
        writer_a.join(timeout=10)
        writer_b.join(timeout=10)

        with sqlite3.connect(str(db_path)) as conn:
            mode = conn.execute("PRAGMA journal_mode").fetchone()[0]

        assert mode == "wal", f"Expected WAL mode but got {mode!r}"
