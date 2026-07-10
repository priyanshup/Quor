"""DOCX structure extraction (QB-007E2).

Converts a Microsoft Word `.docx` file into Markdown-shaped plain text —
headings, paragraphs, bullet/numbered lists, GitHub-style tables, and
contiguous code-style paragraphs as fenced blocks — using the optional
`python-docx` dependency (`quor[documents]`).

`extract_docx()` is fail-open on its own — every failure mode (missing
dependency, corrupt file, invalid zip, unreadable file, unexpected parser
exception) returns `None` and never raises, independent of whatever calls
it. `quor/pipeline/extract/registry.py`'s `extract()` also wraps every
handler in the same guarantee (defense in depth, and the only thing the
still-unimplemented `.pdf` stub relies on), so this is deliberate,
non-load-bearing duplication, not an oversight.

Deliberately excluded by construction, not by filtering after the fact:
document properties (`document.core_properties` — author, revision,
timestamps), comments, and headers/footers are never read at all, because
only `document.element.body`'s paragraphs and tables are walked.
"""

from __future__ import annotations

import re
import warnings
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

# Heading style names python-docx exposes for Word's built-in Heading 1..6
# styles. Heading 7-9 exist in Word but are out of this phase's stated scope
# (ATX Markdown headings only go to ######) and are treated as normal body
# paragraphs, same as "Title"/"Subtitle" — not silently downgraded further,
# just not specially recognized.
_HEADING_STYLE_RE = re.compile(r"^Heading ([1-6])$")

# Style-name substring heuristic for code-style paragraphs (e.g. a "Code" or
# "HTML Code" style many templates define). Checked case-insensitively.
_CODE_STYLE_KEYWORD = "code"

# Explicit run-level font override heuristic for code-style paragraphs, used
# when the paragraph's *style* doesn't already say "code" (e.g. a "Normal"
# paragraph where the author just switched the font to a monospace face).
# Only an explicit per-run override is detected — a monospace convention
# expressed purely via a custom style or theme font is not (see module
# docstring's limitations note in the QB-007E2 backlog entry).
_MONOSPACE_FONTS = frozenset(
    {
        "consolas",
        "courier new",
        "courier",
        "monaco",
        "menlo",
        "source code pro",
        "fira code",
        "cascadia code",
        "cascadia mono",
        "lucida console",
    }
)


def extract_docx(file_path: Path) -> str | None:
    """Return Markdown-shaped plain text for a `.docx` file, or `None`.

    `None` covers two distinct cases, both fail-open and both logged with a
    warning rather than raised:
      - `python-docx` is not installed — a specific, actionable message
        (install `quor[documents]`), checked first so it's never masked by
        the generic one below.
      - anything else went wrong opening or rendering the file — corrupt
        file, invalid zip, unreadable/missing file, or any other unexpected
        exception from python-docx/lxml/zipfile. This function never raises.
    """
    try:
        import docx
    except ImportError:
        warnings.warn(
            "[quor] python-docx is not installed; install quor[documents] "
            "to enable DOCX extraction",
            stacklevel=2,
        )
        return None

    try:
        document = docx.Document(str(file_path))
        return _render(document)
    except Exception as exc:  # noqa: BLE001 — fail-open: this handler must never raise on its own
        warnings.warn(f"[quor] document extraction error: {exc}", stacklevel=2)
        return None


def _render(document: Document) -> str:
    """Render a python-docx `Document` to Markdown-shaped text, in the
    document's own paragraph/table order."""
    from docx.table import Table

    blocks: list[str] = []
    code_lines: list[str] = []
    list_counter = 0

    def flush_code() -> None:
        if code_lines:
            blocks.append("```\n" + "\n".join(code_lines) + "\n```")
            code_lines.clear()

    for item in _iter_block_items(document):
        if isinstance(item, Table):
            flush_code()
            list_counter = 0
            rendered = _render_table(item)
            if rendered:
                blocks.append(rendered)
            continue

        paragraph = item
        raw_text = paragraph.text
        text = raw_text.strip()
        if not text:
            # A blank paragraph ends any code/list run in progress but is
            # otherwise not itself rendered as an empty Markdown block.
            flush_code()
            list_counter = 0
            continue

        style_name = _style_name(paragraph)
        heading_match = _HEADING_STYLE_RE.match(style_name)

        if heading_match:
            flush_code()
            list_counter = 0
            level = int(heading_match.group(1))
            blocks.append(f"{'#' * level} {_flatten(text)}")
        elif style_name.startswith("List Bullet"):
            flush_code()
            list_counter = 0
            blocks.append(f"- {_flatten(text)}")
        elif style_name.startswith("List Number"):
            flush_code()
            list_counter += 1
            blocks.append(f"{list_counter}. {_flatten(text)}")
        elif _is_code_paragraph(paragraph, style_name):
            # Only trailing whitespace is trimmed here — unlike every other
            # branch, leading whitespace is significant (indentation) inside
            # a fenced code block and must survive verbatim.
            list_counter = 0
            code_lines.append(raw_text.rstrip())
        else:
            flush_code()
            list_counter = 0
            blocks.append(text)

    flush_code()
    return "\n\n".join(blocks)


def _iter_block_items(document: Document) -> list[Paragraph | Table]:
    """Yield each paragraph and table in `document.element.body`, in the
    order they actually appear — `document.paragraphs`/`document.tables`
    are separate flat lists that lose this interleaving entirely. This is
    python-docx's own documented recipe for in-order iteration (there is no
    single built-in property for it)."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    items: list[Paragraph | Table] = []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            items.append(Paragraph(child, document))
        elif child.tag == qn("w:tbl"):
            items.append(Table(child, document))
    return items


def _style_name(paragraph: Paragraph) -> str:
    style = paragraph.style
    return (getattr(style, "name", "") or "") if style is not None else ""


def _is_code_paragraph(paragraph: Paragraph, style_name: str) -> bool:
    if _CODE_STYLE_KEYWORD in style_name.lower():
        return True
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all((r.font.name or "").strip().lower() in _MONOSPACE_FONTS for r in runs)


def _render_table(table: Table) -> str:
    """Render a python-docx `Table` as a GitHub-style Markdown table. The
    first row is always treated as the header — python-docx has no general
    way to detect "is this actually a header row," and GitHub-style tables
    require one. A horizontally-merged cell repeats the same text across
    every grid column it spans (Markdown tables have no colspan syntax to
    represent a true merge)."""
    rows = [[_flatten_cell(cell) for cell in row.cells] for row in table.rows]
    if not rows:
        return ""

    col_count = len(rows[0])
    if col_count == 0:
        return ""

    lines = [
        "| " + " | ".join(rows[0]) + " |",
        "| " + " | ".join(["---"] * col_count) + " |",
    ]
    for row in rows[1:]:
        cells = (row + [""] * col_count)[:col_count]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _flatten_cell(cell: object) -> str:
    """A table cell can contain multiple paragraphs; join them with `<br>`
    (the standard GitHub Markdown convention) rather than a space, so
    distinct lines aren't silently merged into one — and escape `|` so
    cell content can never be mistaken for a column boundary."""
    text = "\n".join(p.text for p in cell.paragraphs).strip()  # type: ignore[attr-defined]
    text = text.replace("|", "\\|")
    return text.replace("\n", "<br>")


def _flatten(text: str) -> str:
    """Collapse internal whitespace (including python-docx's own `\\t`/`\\n`
    mappings for tabs/line-breaks) to single spaces, so headings and list
    items — which must stay on one line to match the existing `markdown`
    filter's single-line `preserve_patterns` — can never accidentally span
    multiple lines."""
    return " ".join(text.split())
